"""
rag/templates/docx_renderer.py — render a narrative template as .docx.

For the 14 narrative v2 anchors (and any future narrative leaves),
Word is the native format: compliance officers edit prose, track
changes, leave comments. This module converts the rendered markdown
body into a python-docx Document with reasonable Word styling.

Phase A scope (download-only):
  - Convert common markdown idioms to Word-native paragraph styles
    (headings, blockquotes, bullets, bold/italic runs, code spans).
  - Preserve <<MUST item:X>> markers + EDIT-ZONE comments as visible
    plain text so the tenant sees the structural anchors and can
    edit between them. Phase B (round-trip) will move these to
    Word-native comments/bookmarks.
  - <<TEXT>> placeholders render as a styled "Click to enter text"
    cue so the editable region is visually obvious.

Not full markdown — focused on the 80% of constructs the v2 templates
actually use. Tables, footnotes, link references, etc. are
unhandled; narrative templates don't use them.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ── Constants ───────────────────────────────────────────────────────────────

MUST_MARKER_RE   = re.compile(r"<<(MUST|SHOULD)\s+(item:[A-Za-z0-9.]+:[a-z0-9_]+)>>")
EDIT_ZONE_START  = re.compile(r"<!--\s*EDIT-ZONE-START\s+(item:[A-Za-z0-9.]+:[a-z0-9_]+)\s*-->")
EDIT_ZONE_END    = re.compile(r"<!--\s*EDIT-ZONE-END\s+(item:[A-Za-z0-9.]+:[a-z0-9_]+)\s*-->")
HTML_COMMENT_RE  = re.compile(r"^<!--.*?-->\s*$", re.DOTALL)
PROV_COMMENT_RE  = re.compile(r"^<!--.+?-->\n*", re.DOTALL)
# Ship 54'.d — doc-control + revision-history markers. When a template
# includes <<DOC_CONTROL>>, the renderer emits a document-control
# table (Doc No / Rev / Prepared / Reviewed / Approved / Date + wet-
# sign lines). <<REVISION_HISTORY>> emits an empty audit-defensible
# revision history table. Optional per-template — not every template
# is a controlled document.
DOC_CONTROL_RE       = re.compile(r"^\s*<<DOC_CONTROL>>\s*$")
REVISION_HISTORY_RE  = re.compile(r"^\s*<<REVISION_HISTORY>>\s*$")
# Task #603 iteration (2026-08-15) — detect GFM pipe-tables in the
# rendered markdown body and lift them to native Word tables. Header
# row starts with `|`, ends with `|`; the following line is the
# separator (only `|`, `-`, `:`, whitespace). Without this pass the
# doc-control / revision-history tables (which renderer.py already
# substituted from <<DOC_CONTROL>> markers into markdown pipe-tables)
# would land in the docx as literal `| Document No. |` text.
_PIPE_TABLE_ROW_RE       = re.compile(r"^\s*\|.*\|\s*$")
_PIPE_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|[\s:\-|]+\|\s*$")
# Task #604 (2026-08-15) — round-trip binding markers that must survive
# on disk but should not display in Word's normal view. Any paragraph
# containing one of these gets its runs marked `w:vanish`.
_BINDING_MARKER_RE       = re.compile(r"<<(?:MUST|SHOULD)\s+item:")


def _hide_runs(paragraph) -> None:
    """Task #604 (2026-08-15) — flag every run in the paragraph as
    hidden text (OOXML `w:vanish`). The content survives on disk for
    round-trip extractor binding; Word doesn't display it unless the
    reader toggles \"Show/hide formatting marks\" or \"Hidden text\"."""
    for run in paragraph.runs:
        rpr = run._element.get_or_add_rPr()
        # If an earlier pass already stamped vanish, don't double-stamp.
        if rpr.find(qn("w:vanish")) is None:
            rpr.append(OxmlElement("w:vanish"))


def _add_permission_range(paragraph, kind: str, perm_id: int) -> None:
    """Task #604 (2026-08-15) — insert a `w:permStart` / `w:permEnd`
    marker into the paragraph. `kind` is 'start' or 'end'. Under
    document protection (`readOnly` + enforcement=1), the range
    between a matching start/end pair with `w:edGrp="everyone"`
    remains editable while the rest of the doc is locked."""
    tag = "w:permStart" if kind == "start" else "w:permEnd"
    elem = OxmlElement(tag)
    elem.set(qn("w:id"), str(perm_id))
    if kind == "start":
        elem.set(qn("w:edGrp"), "everyone")
    # Insert as the first child of the paragraph body — permStart
    # applies from that point on; permEnd closes it.
    paragraph._p.insert(0, elem)


def _enable_document_protection(doc) -> None:
    """Task #604 (2026-08-15) — set `w:documentProtection` in
    settings.xml so the whole document is read-only except for the
    permission ranges stamped around edit zones. Enforcement is
    unpasswordded (`w:enforcement=1` alone; no `w:cryptProviderType`)
    so tenants can save their fills without a password prompt but
    Word still guards the scaffolding."""
    settings = doc.settings.element
    # Remove any prior documentProtection to keep this idempotent
    for prior in settings.findall(qn("w:documentProtection")):
        settings.remove(prior)
    prot = OxmlElement("w:documentProtection")
    prot.set(qn("w:edit"),        "readOnly")
    prot.set(qn("w:enforcement"), "1")
    settings.append(prot)
# Ship 57' — per-leaf prerequisites marker. Sits once per template
# (typically under "Before you start"); rendered as a grouped callout
# based on the target leaf_id. Empty prereqs ⇒ marker silently dropped.
PREREQUISITES_MARKER_RE = re.compile(r"^\s*<<PREREQUISITES>>\s*$")

# Templates-Pass-1 (2026-08-08) — <<CROSS_REFERENCES>> marker. NO-OP
# placeholder until Pass 4 wires the Neo4j xfw-bridge resolver. Currently
# consumed silently so Pass-1-normalized templates render cleanly.
CROSS_REFERENCES_MARKER_RE = re.compile(r"^\s*<<CROSS_REFERENCES>>\s*$")

# Ship 56'.a — per-MUST guidance marker. Interleaved between the
# MUST/SHOULD marker and the tenant's <<TEXT>> placeholder; resolves to
# the guidance array on the preceding ChecklistItem.
GUIDANCE_MARKER_RE  = re.compile(r"^\s*<<GUIDANCE>>\s*$")
INLINE_RUN_RE    = re.compile(
    r"(\*\*[^*\n]+\*\*"     # **bold**
    r"|__[^_\n]+__"          # __bold__
    r"|\*[^*\n]+\*"          # *italic*
    r"|_[^_\n]+_"            # _italic_
    r"|`[^`\n]+`"            # `code`
    r")"
)


# ── Inline run helpers ──────────────────────────────────────────────────────

def _add_runs_with_formatting(paragraph, text: str, base_color: RGBColor | None = None) -> None:
    """Split text on inline-markdown boundaries and add runs with bold /
    italic / monospace styling per fragment."""
    if not text:
        return
    parts = INLINE_RUN_RE.split(text)
    for part in parts:
        if not part:
            continue
        bold = italic = code = False
        content = part
        if (part.startswith("**") and part.endswith("**")) or \
           (part.startswith("__") and part.endswith("__")):
            bold = True
            content = part[2:-2]
        elif (part.startswith("*") and part.endswith("*") and len(part) > 2):
            italic = True
            content = part[1:-1]
        elif (part.startswith("_") and part.endswith("_") and len(part) > 2):
            italic = True
            content = part[1:-1]
        elif part.startswith("`") and part.endswith("`"):
            code = True
            content = part[1:-1]
        run = paragraph.add_run(content)
        run.bold   = bold
        run.italic = italic
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        if base_color:
            run.font.color.rgb = base_color


def _add_paragraph(doc, text: str, style: str | None = None,
                   alignment=None, indent_cm: float | None = None,
                   color: RGBColor | None = None):
    """Add a styled paragraph; inline markdown handled.
    Returns the resulting paragraph so callers (e.g. Task #604's
    binding-marker hider) can post-process it."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    _add_runs_with_formatting(p, text, base_color=color)
    return p


# ── Marker-aware line processing ────────────────────────────────────────────

def _humanize_item_slug(mid: str) -> str:
    """`item:4.3:boundaries` → 'boundaries'.

    Drops the internal `item:CTRL:` prefix so the tenant-visible marker
    reads as a natural label ('boundaries', 'exclusions', 'owner')
    rather than a debug-log tag. The full id remains encoded in the
    round-trip `<<MUST item:X>>` marker preserved in the source .md,
    just not visually loud in the rendered doc."""
    if not mid:
        return ""
    from rag.id_types import item_slug
    slug = item_slug(mid) or mid
    return slug.replace("_", " ")


def _render_marker_line(doc, marker_match: re.Match) -> None:
    """Render a <<MUST item:X>> or <<SHOULD item:X>> marker as a small
    label paragraph — visible structural anchor."""
    kind, mid = marker_match.group(1), marker_match.group(2)
    kind_label = "Required element" if kind == "MUST" else "Recommended addition"
    slug = _humanize_item_slug(mid)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"◆ {kind_label} — {slug}")
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x1B)


def _render_edit_zone_marker(
    doc, label: str, mid: str, perm_id: int | None = None,
) -> None:
    """Render an EDIT-ZONE-START / EDIT-ZONE-END line as a subtle
    guidance cue. The internal `item:X:Y` id is dropped from the
    visible text — the surrounding structural markers keep round-trip
    binding.

    Task #604 (2026-08-15) — when `perm_id` is supplied, stamp a
    `w:permStart` on the ▽ paragraph and a `w:permEnd` on the △
    paragraph with matching id + `w:edGrp="everyone"`. Under
    `w:documentProtection readOnly enforcement=1` (enabled at the
    end of the walker), the range between them stays editable
    while the rest of the doc is locked.
    """
    slug = _humanize_item_slug(mid)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    text = (f"▽ Enter your evidence for “{slug}” below ▽"
            if label == "EDIT START"
            else f"△ End of “{slug}” △")
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0xBA, 0xB8, 0xAB)
    # Task #604 (2026-08-15) — stamp permStart AFTER the ▽ rail
    # paragraph and permEnd BEFORE the △ rail paragraph, at body level.
    # This puts BOTH the rails themselves under document protection
    # while the paragraph(s) strictly between them stay editable —
    # so the tenant can add multiple lines of evidence with Enter,
    # and can't accidentally delete the ▽/△ rails (which the
    # extractor uses to find the edit zone).
    if perm_id is not None:
        elem_tag = "w:permStart" if label == "EDIT START" else "w:permEnd"
        elem = OxmlElement(elem_tag)
        elem.set(qn("w:id"), str(perm_id))
        if label == "EDIT START":
            elem.set(qn("w:edGrp"), "everyone")
            # Insert AFTER this ▽ paragraph as a body-level sibling.
            p._p.addnext(elem)
        else:
            # Insert BEFORE the △ paragraph as a body-level sibling.
            p._p.addprevious(elem)


def _render_guidance_block(doc, guidance: tuple[str, ...] | list[str]) -> None:
    """Render a per-MUST guidance callout: bold 'Best practice:' label
    followed by an indented bullet list of imperative steps. Ship 56'.a."""
    if not guidance:
        return
    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_before = Pt(6)
    label_p.paragraph_format.space_after  = Pt(2)
    lr = label_p.add_run("Best practice:")
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    for g in guidance:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent   = Cm(0.8)
        bp.paragraph_format.space_before  = Pt(0)
        bp.paragraph_format.space_after   = Pt(2)
        _add_runs_with_formatting(bp, g)
        for run in bp.runs:
            if run.font.size is None:
                run.font.size = Pt(10)


def _render_prerequisites_block(doc, prereqs) -> None:
    """Render the per-leaf prerequisites callout: bold "Prerequisites:"
    label, then grouped entries (foundational → direct → cross_role)
    each with a ref+title heading, a Why line, and an optional
    Good-enough line. Ship 57'."""
    if not prereqs:
        return
    order = ("foundational", "direct", "cross_role")
    by_cat: dict[str, list] = {}
    for p in prereqs:
        by_cat.setdefault(p.category, []).append(p)

    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_before = Pt(6)
    label_p.paragraph_format.space_after  = Pt(2)
    lr = label_p.add_run("Prerequisites:")
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    for cat in order:
        for p in by_cat.get(cat, []):
            ref_display = _humanize_std_ref(p.standard_id, p.ref)
            head = doc.add_paragraph(style="List Bullet")
            head.paragraph_format.left_indent  = Cm(0.8)
            head.paragraph_format.space_before = Pt(2)
            head.paragraph_format.space_after  = Pt(0)
            head_run = head.add_run(f"{ref_display} — {p.title}")
            head_run.bold = True
            head_run.font.size = Pt(10)

            why = doc.add_paragraph()
            why.paragraph_format.left_indent  = Cm(1.4)
            why.paragraph_format.space_before = Pt(0)
            why.paragraph_format.space_after  = Pt(0)
            _add_runs_with_formatting(why, f"Why: {p.rationale}")
            for run in why.runs:
                if run.font.size is None:
                    run.font.size = Pt(10)

            if p.good_enough:
                ge = doc.add_paragraph()
                ge.paragraph_format.left_indent  = Cm(1.4)
                ge.paragraph_format.space_before = Pt(0)
                ge.paragraph_format.space_after  = Pt(2)
                _add_runs_with_formatting(ge, f"Good enough: {p.good_enough}")
                for run in ge.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)


def _render_cross_references_block(doc, bridges) -> None:
    """Render the outbound xfw-bridge callout: bold "Cross-references:"
    label, then groups by edge type (IMPLEMENTS/SUPPORTS/ENABLES/GOVERNANCE),
    each with target ref+title + curator rationale. Templates Pass 4."""
    if not bridges:
        return
    order = ("IMPLEMENTS", "SUPPORTS", "ENABLES", "GOVERNANCE")
    labels = {
        "IMPLEMENTS": "Implements:",
        "SUPPORTS":   "Supports:",
        "ENABLES":    "Enables:",
        "GOVERNANCE": "Provides governance for:",
    }
    by_type: dict[str, list] = {}
    for b in bridges:
        by_type.setdefault(b.edge_type, []).append(b)

    label_p = doc.add_paragraph()
    label_p.paragraph_format.space_before = Pt(6)
    label_p.paragraph_format.space_after  = Pt(2)
    lr = label_p.add_run("Cross-references:")
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    for et in order:
        entries = by_type.get(et, [])
        if not entries:
            continue
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.left_indent   = Cm(0.4)
        sub_p.paragraph_format.space_before  = Pt(4)
        sub_p.paragraph_format.space_after   = Pt(1)
        sub_run = sub_p.add_run(labels[et])
        sub_run.italic = True
        sub_run.font.size = Pt(10)

        for b in entries:
            std_ref = _humanize_std_ref(b.dst_std, b.dst_ref)
            title = f" — {b.dst_title}" if b.dst_title else ""
            head = doc.add_paragraph(style="List Bullet")
            head.paragraph_format.left_indent  = Cm(0.8)
            head.paragraph_format.space_before = Pt(1)
            head.paragraph_format.space_after  = Pt(0)
            head_run = head.add_run(f"{std_ref}{title}")
            head_run.bold = True
            head_run.font.size = Pt(10)

            if b.rationale:
                r_p = doc.add_paragraph()
                r_p.paragraph_format.left_indent  = Cm(1.4)
                r_p.paragraph_format.space_before = Pt(0)
                r_p.paragraph_format.space_after  = Pt(2)
                _add_runs_with_formatting(r_p, b.rationale)
                for run in r_p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)


def _humanize_std_ref(standard_id: str, ref: str) -> str:
    """Compact display form for prereq ref+std headers."""
    if standard_id.startswith("GDPR:"):
        return f"GDPR {ref}"
    if standard_id.startswith("ISO27001:"):
        return f"ISO 27001 {ref}"
    if standard_id.startswith("ISO27701:"):
        return f"ISO 27701 {ref}"
    return f"{standard_id} {ref}"


def _render_text_placeholder(doc) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("[ Click to enter your evidence here ]")
    run.italic         = True
    run.font.size      = Pt(11)
    run.font.color.rgb = RGBColor(0x9C, 0x9A, 0x8E)


# ── Main converter ──────────────────────────────────────────────────────────

# ── Ship 54'.d — doc-control + revision-history block renderers ─────────

def _derive_doc_number(leaf_id: str, template_version: int | None) -> str:
    """Compose a human-readable document number from the leaf_id.

    Examples:
      req:5.2:information_security_policy → ISP-5.2-Rev00
      req:A.5.15:communication_record     → REC-A.5.15-Rev00
      req:Art.30:records_of_processing    → ROPA-Art.30-Rev00

    Convention: {TYPE_PREFIX}-{control_ref}-Rev{template_version:02d}.
    The TYPE_PREFIX is derived from the leaf_type token (e.g. policy →
    POL, procedure → PRC, register → REG, record → REC). Falls back
    to the raw leaf_type slug when unknown so the doc number is always
    populated + curator-overridable in future via template YAML.
    """
    parts = (leaf_id or "").split(":")
    if len(parts) < 3:
        return leaf_id or "DOC"
    ctrl_ref, leaf_type = parts[1], parts[2]
    lt = leaf_type.lower()
    if "policy" in lt:
        prefix = "POL"
    elif "procedure" in lt or "process" in lt:
        prefix = "PRC"
    elif "register" in lt:
        prefix = "REG"
    elif "record" in lt or "log" in lt:
        prefix = "REC"
    elif "review" in lt:
        prefix = "REV"
    elif "framework" in lt:
        prefix = "FRM"
    elif "manual" in lt:
        prefix = "MAN"
    elif "scope" in lt:
        prefix = "SCP"
    else:
        prefix = "DOC"
    rev = f"Rev{(template_version or 1):02d}"
    return f"{prefix}-{ctrl_ref}-{rev}"


def _render_doc_control_block(
    doc, leaf_id: str, template_version: int | None,
) -> None:
    """Insert a document-control table matching consultant-toolkit
    convention (mirrors Share.zip L2-PRC-003 doc-control block).

    Layout (2-column table):
      Document No.       | {derived from leaf_id}
      Revision           | Rev{template_version:02d}
      Revision Date      | {today, DD MMM YYYY}
      Prepared By        | ___________________________ (wet-sign)
      Reviewed By        | ___________________________
      Approved By        | ___________________________

    Deliberate blanks on Prepared/Reviewed/Approved — the tenant fills
    those at document-control review. No auto-fill from tenant_profile
    for the MVP (deferred; tenant_profile doesn't carry
    prepared_by / reviewed_by fields).
    """
    from datetime import date
    doc_no = _derive_doc_number(leaf_id, template_version)
    today  = date.today().strftime("%d %b %Y")
    rev    = f"Rev{(template_version or 1):02d}"

    rows: list[tuple[str, str]] = [
        ("Document No.",   doc_no),
        ("Revision",       rev),
        ("Revision Date",  today),
        ("Prepared By",    "___________________________"),
        ("Reviewed By",    "___________________________"),
        ("Approved By",    "___________________________"),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, (label, value) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.size = Pt(10)

    # Small trailing paragraph to separate from body content
    doc.add_paragraph()


def _render_revision_history_block(
    doc, template_version: int | None,
) -> None:
    """Insert a revision-history table. Empty audit-defensible shape:
    curator/tenant adds one row per version. Convention matches Share
    L2-PRC-003 revision-history block.

    Columns: Version | Date | Description of Change | Author

    Seeded with one row for the current template_version + today's
    date. Future edits append rows.
    """
    from datetime import date

    doc.add_heading("Revision History", level=2)

    table = doc.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    # Header row
    hdr = ["Version", "Date", "Description of Change", "Author"]
    for i, h in enumerate(hdr):
        c = table.cell(0, i)
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)

    # Current-version seed row
    today = date.today().strftime("%d %b %Y")
    rev = f"{(template_version or 1):02d}"
    seed = [rev, today, "Initial issue / current version", ""]
    for i, v in enumerate(seed):
        c = table.cell(1, i)
        c.text = ""
        r = c.paragraphs[0].add_run(v)
        r.font.size = Pt(10)

    doc.add_paragraph()


def _split_pipe_row(line: str) -> list[str]:
    """Split a GFM pipe-table row into stripped cell values.
    Leading/trailing pipes discarded; empty cells preserved."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _extract_pipe_tables(body: str) -> tuple[str, list[dict]]:
    """Task #603 (2026-08-15) — pre-pass over the markdown body: find
    every GFM pipe-table (header row + separator + zero-or-more data
    rows) and replace it with a single sentinel line `<<TABLE::N>>`.
    Returns the rewritten body plus the parsed table specs.

    Each table spec: {"headers": [str, ...], "rows": [[str, ...], ...]}.

    Sentinel dispatching happens in the main line walker below —
    _render_pipe_table lifts the spec into a native Word table.
    """
    lines = body.splitlines(keepends=True)
    out_parts: list[str] = []
    tables: list[dict] = []
    i = 0
    while i < len(lines):
        # A pipe-table starts with a header row followed by a separator row.
        if (i + 1 < len(lines)
                and _PIPE_TABLE_ROW_RE.match(lines[i])
                and _PIPE_TABLE_SEPARATOR_RE.match(lines[i + 1])):
            headers = _split_pipe_row(lines[i])
            data_rows: list[list[str]] = []
            j = i + 2
            while j < len(lines) and _PIPE_TABLE_ROW_RE.match(lines[j]):
                data_rows.append(_split_pipe_row(lines[j]))
                j += 1
            table_idx = len(tables)
            tables.append({"headers": headers, "rows": data_rows})
            out_parts.append(f"<<TABLE::{table_idx}>>\n")
            i = j
            continue
        out_parts.append(lines[i])
        i += 1
    return "".join(out_parts), tables


_TABLE_SENTINEL_RE = re.compile(r"^\s*<<TABLE::(\d+)>>\s*$")


_SIGNATURE_CELL_RE = re.compile(r"^_{3,}$")


def _render_pipe_table(doc, spec: dict, perm_counter: list[int] | None = None) -> None:
    """Emit a parsed pipe-table as a native python-docx table. Header
    row bolded; data cells preserve markdown formatting (bold/italic
    runs inside cells) via `_add_runs_with_formatting`. Empty tables
    still emit the header for auditor-visible structure.

    Task #605 (2026-08-15) — cells whose value matches the wet-sign
    pattern (`^_{3,}$` after stripping) get wrapped with matching
    `w:permStart` + `w:permEnd` at cell-body level. Under document
    protection those cells stay editable while the rest of the doc
    is locked. `perm_counter` is a mutable box [int]; if None,
    signature-cell editability is skipped (i.e. rendered as fixed
    scaffolding). Closes docx lockdown dogfood friction #2 —
    Prepared By / Reviewed By / Approved By.
    """
    headers = spec.get("headers") or []
    rows    = spec.get("rows") or []
    if not headers:
        return
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    # Header row — force bold even when the source markdown lacked **…**.
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        cell.text = ""
        p = cell.paragraphs[0]
        _add_runs_with_formatting(p, h)
        for run in p.runs:
            run.bold = True
            if run.font.size is None:
                run.font.size = Pt(10)
    # Data rows — respect cell-level markdown (**bold**, _italic_, `code`).
    for r_idx, row in enumerate(rows, start=1):
        for c in range(n_cols):
            cell = table.cell(r_idx, c)
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[c] if c < len(row) else ""
            _add_runs_with_formatting(p, val)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(10)
            # Task #605 — wet-sign cell → permStart/permEnd around the
            # cell body so the wet-sign underscores stay typable under
            # document protection.
            if perm_counter is not None and _SIGNATURE_CELL_RE.match(val.strip()):
                perm_counter[0] += 1
                perm_id = perm_counter[0]
                tc = cell._tc  # the underlying w:tc element
                ps = OxmlElement("w:permStart")
                ps.set(qn("w:id"),    str(perm_id))
                ps.set(qn("w:edGrp"), "everyone")
                pe = OxmlElement("w:permEnd")
                pe.set(qn("w:id"),    str(perm_id))
                # Insert permStart as the first cell child (before all
                # paragraphs) and permEnd as the last cell child. Word
                # treats the range as covering the whole cell body.
                first_child = tc[0] if len(tc) else None
                if first_child is not None:
                    first_child.addprevious(ps)
                else:
                    tc.append(ps)
                tc.append(pe)
    doc.add_paragraph()  # separator after the table


def render_template_docx(
    pg_conn,
    tenant_id: str,
    leaf_id:   str,
    template_body: str,
    template_version: int | None = None,
) -> bytes:
    """Convert the rendered markdown body to a .docx workbook.

    Walks line-by-line, maps markdown idioms to Word styles, and
    preserves structural markers as visible cues. Returns bytes
    suitable for HTTP response.

    Ship 54'.d — the walk also detects <<DOC_CONTROL>> and
    <<REVISION_HISTORY>> markers and dispatches to the block
    renderers. `template_version` is passed through so those blocks
    can render the current Rev number.
    """
    doc = Document()

    # Set base style — Word's default Calibri 11pt is fine; tighten margins
    # slightly so compliance officers don't waste a page on whitespace.
    sec = doc.sections[0]
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    # Strip the leading provenance HTML comment (the render_template
    # `include_header=True` provenance line — it's metadata, not content).
    body = PROV_COMMENT_RE.sub("", template_body, count=1)

    # Task #603 (2026-08-15) — pre-pass extracts GFM pipe-tables and
    # replaces them with `<<TABLE::N>>` sentinels; the parsed specs
    # are lifted to native Word tables inside the main walker below.
    body, _pipe_tables = _extract_pipe_tables(body)

    in_code_block = False
    in_list_block = False
    current_item_id: str | None = None   # Ship 56'.a — set on MUST/SHOULD marker;
                                          # consumed by <<GUIDANCE>> marker.
    # Task #604 (2026-08-15) — permission-range counter. Each edit zone
    # (EDIT-ZONE-START / END pair) gets a unique w:permStart/w:permEnd
    # pair id. After the walk finishes we set w:documentProtection so
    # the rest of the doc is locked.
    _perm_counter = [0]
    _perm_open_id: int | None = None

    for raw_line in body.splitlines():
        line = raw_line.rstrip()

        # Code fences — flip state, render a horizontal-rule-ish separator
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Empty line — paragraph break
        if not line.strip():
            in_list_block = False
            continue

        # Task #603 (2026-08-15) — sentinel for a pipe-table lifted by
        # the pre-pass. Dispatch to native Word table renderer.
        # Task #605 — thread the perm counter so wet-sign cells inside
        # the doc-control table become editable ranges.
        _t = _TABLE_SENTINEL_RE.match(line)
        if _t:
            _spec = _pipe_tables[int(_t.group(1))]
            _render_pipe_table(doc, _spec, perm_counter=_perm_counter)
            continue

        # Ship 54'.d — <<DOC_CONTROL>> marker → document-control table
        if DOC_CONTROL_RE.match(line):
            _render_doc_control_block(doc, leaf_id, template_version)
            continue

        # Ship 54'.d — <<REVISION_HISTORY>> marker → revision history table
        if REVISION_HISTORY_RE.match(line):
            _render_revision_history_block(doc, template_version)
            continue

        # MUST / SHOULD marker as its own line
        marker = MUST_MARKER_RE.match(line.strip())
        if marker:
            current_item_id = marker.group(2)
            _render_marker_line(doc, marker)
            continue

        # Ship 56'.a — <<GUIDANCE>> marker → per-MUST guidance callout
        # for the preceding MUST/SHOULD's ChecklistItem. Empty list ⇒
        # marker silently dropped.
        if GUIDANCE_MARKER_RE.match(line):
            if current_item_id:
                from rag.templates.guidance_lookup import get_guidance_for_item
                _render_guidance_block(doc, get_guidance_for_item(current_item_id))
            continue

        # Ship 57' — <<PREREQUISITES>> marker → per-leaf prereqs callout.
        # Uses the template's target leaf_id (parameter to this function),
        # not a preceding item marker. Empty prereqs ⇒ marker dropped.
        if PREREQUISITES_MARKER_RE.match(line):
            from rag.templates.prerequisites_lookup import get_prerequisites_for_leaf
            _render_prerequisites_block(doc, get_prerequisites_for_leaf(leaf_id))
            continue

        # Templates Pass 4 (2026-08-08) — <<CROSS_REFERENCES>> marker →
        # outbound xfw-bridge callout for the leaf's control. Empty
        # bridges ⇒ marker dropped.
        if CROSS_REFERENCES_MARKER_RE.match(line):
            from rag.templates.cross_references_lookup import get_cross_references_for_leaf
            _render_cross_references_block(doc, get_cross_references_for_leaf(leaf_id))
            continue

        # EDIT-ZONE markers. Task #604 — stamp matching w:permStart /
        # w:permEnd around each zone using an id from _perm_counter.
        ez_start = EDIT_ZONE_START.match(line.strip())
        if ez_start:
            _perm_counter[0] += 1
            _perm_open_id = _perm_counter[0]
            _render_edit_zone_marker(
                doc, "EDIT START", ez_start.group(1), perm_id=_perm_open_id,
            )
            continue
        ez_end = EDIT_ZONE_END.match(line.strip())
        if ez_end:
            _render_edit_zone_marker(
                doc, "EDIT END",   ez_end.group(1), perm_id=_perm_open_id,
            )
            _perm_open_id = None
            continue

        # <<TEXT>> placeholder
        if line.strip() == "<<TEXT>>":
            _render_text_placeholder(doc)
            continue

        # Other HTML comments — strip silently (provenance, leftover markers)
        if HTML_COMMENT_RE.match(line.strip()):
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}\s*", line.strip()):
            p = doc.add_paragraph()
            run = p.add_run("─" * 40)   # box-drawing horizontal
            run.font.color.rgb = RGBColor(0xC8, 0xC5, 0xB8)
            continue

        # Blockquote (standard-text + advisory blocks)
        if line.startswith("> "):
            content = line[2:].strip()
            p = doc.add_paragraph(style="Intense Quote") if "_Standard text:_" in content \
                else doc.add_paragraph(style="Quote")
            _add_runs_with_formatting(p, content)
            # Task #604 (2026-08-15) — the N/A section wraps the
            # "Do not edit — system id: <<MUST item:X>>" line in a
            # blockquote; catch it here too so the marker stays hidden.
            if _BINDING_MARKER_RE.search(content):
                _hide_runs(p)
            continue
        if line.strip() == ">":
            doc.add_paragraph()
            continue

        # Checklist line (- [ ] or - [x]) — MUST match before generic list
        # else the leading "- " gets eaten by the unordered-list rule and
        # the "[ ]" leaks into the bullet text.
        m = re.match(r"^(\s*)-\s+\[([ xX])\]\s+(.*)$", line)
        if m:
            checked = m.group(2).lower() == "x"
            p = doc.add_paragraph()
            indent_level = len(m.group(1)) // 2
            if indent_level:
                p.paragraph_format.left_indent = Cm(0.6 * (indent_level + 1))
            cb = p.add_run("☒ " if checked else "☐ ")
            cb.font.size = Pt(11)
            _add_runs_with_formatting(p, m.group(3))
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent_level = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            if indent_level:
                p.paragraph_format.left_indent = Cm(0.6 * (indent_level + 1))
            _add_runs_with_formatting(p, m.group(2))
            in_list_block = True
            continue

        # Ordered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_formatting(p, m.group(2))
            in_list_block = True
            continue

        # Default — regular paragraph with inline formatting.
        # Task #604 (2026-08-15) — if the paragraph carries a round-trip
        # binding marker (`<<MUST item:X>>` or `<<SHOULD item:X>>`),
        # emit it as hidden text so the extractor still sees it on
        # re-upload but Word doesn't display it. This is the fix for
        # the dogfood friction "compliance officer sees `<<...>>`".
        p = _add_paragraph(doc, line)
        if _BINDING_MARKER_RE.search(line) and p is not None:
            _hide_runs(p)

    # Task #604 (2026-08-15) — turn on document protection AFTER all
    # paragraphs land (permStart/permEnd markers already stamped inline
    # by the EDIT-ZONE handlers). readOnly + enforcement=1 locks every
    # non-permitted range; the permission ranges around edit zones stay
    # typable. No password — tenants can Save-As freely.
    if _perm_counter[0] > 0:
        _enable_document_protection(doc)

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
