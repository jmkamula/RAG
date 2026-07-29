"""
ArionComply — Document Readers
Stage 1: Extract raw text and structure from uploaded files.
No LLM, no interpretation — pure extraction.

Supported formats:
  PDF   → pdfplumber (text + page markers)
  DOCX  → python-docx (text + heading structure)
  XLSX  → openpyxl (sheets + rows)
  TXT   → direct read
  CSV   → csv module

Dependencies:
  pip install pdfplumber python-docx openpyxl
"""
from __future__ import annotations

import csv
import io
import logging
import os
from pathlib import Path
from typing import Optional

from .models import ParsedDocument, RawSection

logger = logging.getLogger(__name__)

# Token estimate: 1 token ≈ 4 characters (conservative for compliance text)
CHARS_PER_TOKEN = 4


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

def read_document(
    file_path:         str,
    upload_id:         Optional[str] = None,
    original_filename: Optional[str] = None,
) -> ParsedDocument:
    """
    Read a document and return a ParsedDocument with raw sections.
    Dispatches to the appropriate reader based on file extension.

    original_filename: user-facing name (e.g. "Business Continuity Policy.docx").
    Stored on ParsedDocument.original_name and propagated to every Finding's
    document_name, which is what posture_writer uses to match the upload back
    to its pre-registered client_documents row (DOC-prefix / title fuzzy).
    Without this, API uploads (saved as {upload_id}.{ext}) can never link to
    the registry and findings land on a fresh orphan row.
    """
    path      = Path(file_path)
    ext       = path.suffix.lower().lstrip(".")
    file_name = original_filename or path.name

    readers = {
        "pdf":  _read_pdf,
        "docx": _read_docx,
        "doc":  _read_docx,
        "xlsx": _read_xlsx,
        "xlsm": _read_xlsx,
        "xls":  _read_xlsx,
        "txt":  _read_txt,
        "csv":  _read_csv,
        "md":   _read_txt,
    }

    reader = readers.get(ext)
    if reader is None:
        logger.warning(f"Unsupported file type: {ext} — treating as plain text")
        reader = _read_txt

    logger.info(f"Reading {file_name} ({ext})")
    doc = reader(file_path, file_name)
    doc.upload_id = upload_id

    # Compute token estimate. Prefer markdown when present — it's what the
    # extractor actually feeds the LLM, and for table-heavy docs it's
    # materially larger than the paragraph-only join below.
    doc.full_text     = "\n\n".join(s.text for s in doc.raw_sections if s.text.strip())
    sizing_text       = doc.markdown if doc.markdown else doc.full_text
    doc.token_estimate = len(sizing_text) // CHARS_PER_TOKEN

    logger.info(
        f"Read {file_name}: {len(doc.raw_sections)} sections, "
        f"~{doc.token_estimate:,} tokens, {doc.page_count} pages"
    )
    return doc


# =============================================================================
# PDF READER
# =============================================================================

def _read_pdf(file_path: str, file_name: str) -> ParsedDocument:
    """PDF reader — Layer A (2026-06-19): captures tables into a markdown
    rendering alongside the existing section-based text extraction.

    Mirrors the docx → mammoth pattern: extract a markdown representation
    that includes tables (which `page.extract_text()` flattens or drops),
    set on `ParsedDocument.markdown`, and apply the same table-heavy
    rescue + table-prose synthesis as the docx path.

    Closes the 0% bind-rate gap on PDF uploads vs ~14-92% on docx/workbook
    by giving the LLM extractor access to the structural content PDFs
    bury in tables (audit findings, control matrices, certification
    scopes).
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber required: pip install pdfplumber")

    import hashlib

    sections = []
    page_count = 0
    md_parts:  list[str] = []     # per-page markdown blocks
    src_sha:   Optional[str] = None

    # File-level sha (used as src_sha256 telemetry; matches docx pattern)
    try:
        with open(file_path, "rb") as _fh:
            src_sha = hashlib.sha256(_fh.read()).hexdigest()
    except Exception:
        pass

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        current_heading = None
        current_text    = []
        section_start   = 1

        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""

            # --- markdown layer: per-page text + tables ----------------
            page_md: list[str] = []
            if text.strip():
                page_md.append(text.strip())
            try:
                tables = page.extract_tables() or []
            except Exception as e:
                logger.debug(f"pdfplumber table extraction failed on {file_name} page {page_num}: {e}")
                tables = []
            for table in tables:
                md_table = _pdf_table_to_markdown(table)
                if md_table:
                    page_md.append(md_table)
            if page_md:
                md_parts.append(f"## Page {page_num}\n\n" + "\n\n".join(page_md))
            # ---------------------------------------------------------

            if not text.strip():
                continue

            lines = text.splitlines()
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue

                # Detect headings: short lines, possible numbering
                is_heading = (
                    len(stripped) < 80
                    and (
                        # Numbered heading: "1.", "1.1", "4.2 Access Control"
                        bool(__import__('re').match(r'^\d+\.?\d*\s+\w', stripped))
                        # All caps short line
                        or (stripped.isupper() and len(stripped) > 3)
                    )
                )

                if is_heading and current_text:
                    # Save current section
                    sections.append(RawSection(
                        section_id  = f"page_{section_start}_{page_num}",
                        heading     = current_heading,
                        text        = "\n".join(current_text),
                        page_start  = section_start,
                        page_end    = page_num,
                        level       = _detect_heading_level(current_heading or ""),
                    ))
                    current_heading = stripped
                    current_text    = []
                    section_start   = page_num
                elif is_heading:
                    current_heading = stripped
                else:
                    current_text.append(line)

        # Final section
        if current_text:
            sections.append(RawSection(
                section_id  = f"page_{section_start}_{page_count}",
                heading     = current_heading,
                text        = "\n".join(current_text),
                page_start  = section_start,
                page_end    = page_count,
                level       = _detect_heading_level(current_heading or ""),
            ))

    # If no sections detected (flat PDF), treat each page as a section
    if not sections:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    sections.append(RawSection(
                        section_id = f"page_{i}",
                        heading    = f"Page {i}",
                        text       = text,
                        page_start = i,
                        page_end   = i,
                        level      = 0,
                    ))

    md_text: Optional[str] = "\n\n".join(md_parts) if md_parts else None

    # Table-heavy doc rescue (mirrors docx): if paragraph extraction is
    # sparse vs markdown, rebuild sections from markdown so the section-
    # based extractor sees table content.
    para_chars = sum(len(s.text) for s in sections)
    md_chars   = len(md_text or "")
    if md_text and md_chars > max(2000, para_chars * 3):
        sections = _chunk_markdown_to_sections(md_text)
        logger.info(
            f"Rebuilt sections from markdown for {file_name}: "
            f"paragraph_chars={para_chars} md_chars={md_chars} "
            f"-> {len(sections)} chunks (table-heavy PDF rescue)"
        )

    # Table-prose synthesis (mirrors docx): emit per-row sentences so the
    # LLM can cite ≥40-char verbatim quotes from table content.
    table_prose = _synthesise_table_prose(md_text or "")
    if table_prose:
        sections.append(RawSection(
            section_id = "table_prose_synthesis",
            heading    = "Table content (synthesised for extraction)",
            text       = table_prose,
            page_start = None,
            page_end   = None,
            level      = 0,
        ))
        md_text = (md_text or "") + "\n\n" + table_prose
        logger.info(
            f"Table-prose synthesis for {file_name}: emitted "
            f"{len(table_prose)} chars across {table_prose.count('Row ')} rows"
        )

    return ParsedDocument(
        source_file   = file_path,
        file_type     = "pdf",
        original_name = file_name,
        raw_sections  = sections,
        page_count    = page_count,
        markdown      = md_text,
        source_sha256 = src_sha,
        converter     = "pdfplumber",
    )


def _pdf_table_to_markdown(table: list) -> str:
    """Convert a pdfplumber-extracted table (list[list[str|None]]) to
    GitHub-flavoured markdown. Returns "" for empty / single-cell tables
    (not worth the markdown overhead).

    Row 0 is treated as the header — pdfplumber doesn't reliably mark
    headers, but the first row is the convention for most extraction
    consumers. Empty cells (None or whitespace) become a single space
    so the column structure is preserved.
    """
    if not table or len(table) < 2 or not any(table):
        return ""
    # Normalise cells: None → "", strip newlines + outer whitespace, replace
    # pipes (markdown's column separator) with a similar-looking visual.
    def _cell(c) -> str:
        s = (c or "").replace("\n", " ").replace("|", "/").strip()
        return s if s else " "
    rows = [[_cell(c) for c in row] for row in table if row]
    if not rows or all(all(not c.strip() for c in r) for r in rows):
        return ""
    # Pad short rows to the max column count
    ncols = max(len(r) for r in rows)
    rows = [r + [" "] * (ncols - len(r)) for r in rows]
    header = rows[0]
    body   = rows[1:]
    if not body:
        # Single-row "table" — emit as a paragraph not a table
        return " · ".join(c.strip() for c in header if c.strip())
    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


# =============================================================================
# DOCX READER
# =============================================================================

def _read_docx(file_path: str, file_name: str) -> ParsedDocument:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    # Markdown rendering via mammoth — captures tables, headings, and lists
    # that the paragraph walk below ignores. Surfaced on ParsedDocument.markdown
    # and used as the LLM input when present (see extractor._extract_full).
    md_text:     Optional[str] = None
    src_sha:     Optional[str] = None
    md_converter: Optional[str] = None
    try:
        import hashlib, io
        import mammoth
        from importlib.metadata import version as _pkg_version
        with open(file_path, "rb") as _fh:
            _bytes = _fh.read()
        src_sha = hashlib.sha256(_bytes).hexdigest()
        _result = mammoth.convert_to_markdown(io.BytesIO(_bytes))
        md_text = _result.value or None
        try:
            md_converter = f"mammoth/{_pkg_version('mammoth')}"
        except Exception:
            md_converter = "mammoth"

        # Strip embedded base64 image data URIs from the markdown.
        # Mammoth inlines docx images as `![alt](data:image/png;base64,...)`
        # which is pure noise to the LLM (it can't decode base64) and
        # can dominate doc size — observed on 2026-06-10 with ISMS
        # Automation Process.docx: 98.8% of 269KB markdown was a
        # single embedded screenshot. Stripping leaves the alt text +
        # a redaction marker so the LLM still knows an image was
        # present; downstream OCR (future work) can pick up content.
        if md_text:
            md_text = _strip_base64_images(md_text)
            md_text = _strip_word_anchor_tags(md_text)
            # Ship 50'.a — restore templated fast-path for ArionComply-
            # rendered docx uploads. When the ◆ label + attribution
            # signature is present, rewrite the markdown to include
            # `<<MUST item:X:Y>>` markers + edit zones so the extractor's
            # templated path fires. No-op on non-Arion docs.
            md_text = _reconstruct_arion_markers(md_text, file_name)
    except Exception as e:
        logger.warning(f"mammoth markdown conversion failed for {file_name}: {e}")

    doc      = docx.Document(file_path)
    sections = []

    current_heading  = None
    current_level    = 0
    current_paras    = []
    section_idx      = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name

        if is_heading:
            if current_paras:
                sections.append(RawSection(
                    section_id = f"section_{section_idx}",
                    heading    = current_heading,
                    text       = "\n".join(current_paras),
                    page_start = None,
                    page_end   = None,
                    level      = current_level,
                ))
                section_idx += 1

            # Extract heading level from style name (Heading 1, Heading 2, ...)
            import re
            m = re.search(r'(\d+)', style_name)
            current_level   = int(m.group(1)) if m else 1
            current_heading = text
            current_paras   = []
        else:
            current_paras.append(text)

    # Final section
    if current_paras:
        sections.append(RawSection(
            section_id = f"section_{section_idx}",
            heading    = current_heading,
            text       = "\n".join(current_paras),
            page_start = None,
            page_end   = None,
            level      = current_level,
        ))

    # Table-heavy doc rescue: docx tables don't appear in
    # `doc.paragraphs`, so the paragraph walker above misses them
    # entirely. Mammoth's markdown stream captures everything (tables
    # + lists + headings). When paragraph-walk content is small
    # relative to markdown, the doc is table-dominated — rebuild
    # sections from markdown chunks so the section-based extractor
    # actually sees the operative content.
    para_chars = sum(len(s.text) for s in sections)
    md_chars   = len(md_text or "")
    if md_text and md_chars > max(2000, para_chars * 3):
        sections = _chunk_markdown_to_sections(md_text)
        logger.info(
            f"Rebuilt sections from markdown for {file_name}: "
            f"paragraph_chars={para_chars} md_chars={md_chars} "
            f"-> {len(sections)} chunks (table-heavy doc rescue)"
        )

    # Table-prose synthesis: even after the rescue, markdown tables
    # render as `| cell | cell |` rows that the LLM can't cite as
    # ≥40-char verbatim quotes. Synthesise each row as a sentence
    # ("Row N: header1 = cell1; header2 = cell2; ...") and append.
    # Both the LLM input (doc.markdown) and the grounding haystack
    # (doc.full_text — built from sections downstream) pick up the
    # synthesis automatically.
    table_prose = _synthesise_table_prose(md_text or "")
    if table_prose:
        sections.append(RawSection(
            section_id = "table_prose_synthesis",
            heading    = "Table content (synthesised for extraction)",
            text       = table_prose,
            page_start = None,
            page_end   = None,
            level      = 0,
        ))
        md_text = (md_text or "") + "\n\n" + table_prose
        logger.info(
            f"Table-prose synthesis for {file_name}: "
            f"emitted {len(table_prose)} chars across "
            f"{table_prose.count('Row ')} rows"
        )

    return ParsedDocument(
        source_file   = file_path,
        file_type     = "docx",
        original_name = file_name,
        raw_sections  = sections,
        page_count    = 0,  # DOCX doesn't easily expose page count
        markdown      = md_text,
        source_sha256 = src_sha,
        converter     = md_converter,
    )


import re as _re

_BASE64_IMG_RE = _re.compile(
    r"!\[([^\]]*)\]\(data:image/[^)]+\)",
    flags=_re.DOTALL,
)


def _strip_base64_images(md: str) -> str:
    """Replace inline `![alt](data:image/...;base64,...)` references with
    `![image: alt]` markers. Preserves the alt text (which sometimes
    has a useful description) and the structural presence of the image
    so the LLM knows it's there, but drops the unreadable base64 blob
    that bloats the prompt and crowds out real prose.

    Future work: OCR the embedded images (Claude vision or Tesseract)
    and inject the extracted text in place of the marker. For now the
    marker is the placeholder."""
    def _replace(m):
        alt = (m.group(1) or "").strip()
        if alt:
            # Mammoth sometimes embeds a multi-line "AI-generated content
            # may be incorrect" boilerplate in the alt text — keep only
            # the first line for brevity.
            alt = alt.split("\n")[0].strip()
            return f"![image: {alt}]"
        return "![image]"
    return _BASE64_IMG_RE.sub(_replace, md)


# Word-inserted hidden anchor tags — Ship 49'.a.
# Word wraps every heading with `<a id="_Toc..."></a>` (table-of-contents
# targets), `<a id="_Ref..."></a>` (cross-reference targets), and
# `<a id="_Hlk..."></a>` (highlight/comment keys). Mammoth passes them
# through to the markdown stream verbatim. When section headings become
# extractor evidence (Ship 32 multi-attribution), the excerpt reads as
# HTML noise to the reviewer: `# <a id="_Toc..."></a>SCOPE`.
#
# Scrub at read time so downstream stored excerpts are clean. Complements
# the display-side scrubDocxNoise() in static/arioncomply.html — that one
# handles legacy stored data; this prevents future data from being noisy.
_WORD_ANCHOR_TAG_RE = _re.compile(
    r'<a\s+(?:id|name)="_(?:Toc|Ref|Hlk|Bkmk|Anchor)[^"]*"[^>]*>\s*</a>',
    _re.IGNORECASE,
)


def _strip_word_anchor_tags(md: str) -> str:
    """Remove Word-inserted hidden `<a id="_Toc..."></a>` anchor tags
    from a markdown stream. Collapses whitespace runs left behind by
    the removal. Preserves everything else verbatim."""
    scrubbed = _WORD_ANCHOR_TAG_RE.sub("", md)
    # Collapse runs of 2+ spaces (only) — don't touch newline structure.
    scrubbed = _re.sub(r"[ \t]{2,}", " ", scrubbed)
    return scrubbed


# ─── ArionComply template round-trip repair (Ship 50'.a) ─────────────
#
# `rag/templates/docx_renderer.py` transforms `<<MUST item:X:Y>>` markers
# into human-friendly `◆ Required element — <slug>` labels. Beautiful for
# fill-in, but the machine-readable form is lost — on re-upload the
# templated fast-path (`_extract_templated`) can't find markers and
# consensus falls through, matching template scaffolding to random
# controls (Ship 50'.a diagnosis on customer's A_5_1_management_approval).
#
# This block detects an ArionComply-rendered docx from its attribution
# line + ◆ markers, parses the control_ref from the attribution, and
# reconstructs `<<MUST item:CTRL:slug>>` + `<!-- EDIT-ZONE-START/END -->`
# markers so the fast-path fires exactly as if the customer had uploaded
# the source .md.
#
# Attribution line (Quote-styled) in the rendered docx:
#     "Generated 2026-07-29 · A.5.1 · ISO/IEC 27001:2022"
# Mammoth escapes '.' and '-' in markdown output ('A\.5\.1', '2026\-07\-29').
# Allow either escaped or plain form in both the date and the ref.
_ARION_ATTRIBUTION_RE = _re.compile(
    r"Generated\s+\d{4}\\?-\d{2}\\?-\d{2}\s*·\s*([A-Za-z0-9.\\]+?)\s*·",
)
# ◆ Required/Recommended labels — mammoth wraps them in markdown bold
# markers `__...__` since the docx renderer applied Bold to the run.
# Optional surrounding `__` (bold), optional wrapping asterisks, spaces.
_ARION_MUST_LABEL_RE   = _re.compile(
    r"^\s*(?:__|\*\*)?\s*◆\s*Required element\s*—\s*(.+?)\s*(?:__|\*\*)?\s*$",
    _re.MULTILINE,
)
_ARION_SHOULD_LABEL_RE = _re.compile(
    r"^\s*(?:__|\*\*)?\s*◆\s*Recommended addition\s*—\s*(.+?)\s*(?:__|\*\*)?\s*$",
    _re.MULTILINE,
)
# Standalone "Why:" citation line rendered from `_Why: ..._` in the .md.
# Mammoth emits italics as `*...*` — allow optional wrappers.
_ARION_WHY_LABEL_RE = _re.compile(
    r"^\s*\*?Why:\s.+?\*?\s*$", _re.MULTILINE,
)
# "[ Click to enter your evidence here ]" placeholder.
_ARION_CLICK_PLACEHOLDER_RE = _re.compile(
    r"^\s*\[\s*Click to enter your evidence here\s*\]\s*$", _re.MULTILINE,
)


def _detect_arion_control_ref(md_text: str, file_name: str = "") -> str:
    """Return the control_ref for an ArionComply-rendered docx, or ''.

    Primary source: attribution line 'Generated <date> · <control_ref> · <std>'.
    Fallback: filename shape 'A_5_1_<slug>.docx' → 'A.5.1'.
    """
    m = _ARION_ATTRIBUTION_RE.search(md_text)
    if m:
        # Unescape mammoth's `A\.5\.1` → `A.5.1`.
        return m.group(1).replace("\\", "")
    # Filename fallback — strip extension + slug suffix, convert _ → .
    if file_name:
        stem = _re.sub(r"\.\w+$", "", file_name)
        # Split into "A_5_1" (control) + "management_approval" (slug).
        # Convention: uppercase-letter start + underscore-separated
        # number groups is the control part.
        parts = stem.split("_")
        ref_parts: list[str] = []
        for i, p in enumerate(parts):
            if i == 0 and _re.match(r"^[A-Z]$", p):  # 'A' / 'B'
                ref_parts.append(p)
            elif p.isdigit():
                ref_parts.append(p)
            else:
                break
        if len(ref_parts) >= 2:
            return ".".join(ref_parts)
    return ""


def _slug_from_humanized(label: str) -> str:
    """Reverse `_humanize_item_slug` (docx_renderer.py) so 'approval signatory'
    → 'approval_signatory'. Preserves case-safe conversion — the docx
    renderer only replaces `_` with space, so this is exactly the inverse."""
    return label.strip().replace(" ", "_")


def _reconstruct_arion_markers(md_text: str, file_name: str) -> str:
    """When mammoth markdown looks like an ArionComply-rendered docx,
    rewrite it so the templated fast-path can bind evidence per-MUST.

    Detection: ATTRIBUTION line present AND ≥1 ◆ label. Neither alone
    fires — the ◆ pattern could appear in a doc that just happens to
    use that character; the attribution line is our signature.

    Rewrite: each `◆ Required element — <slug>` line is replaced by
    `<<MUST item:CTRL:slug>>`, and an `<!-- EDIT-ZONE-START -->` /
    `<!-- EDIT-ZONE-END -->` bracket is inserted around the paragraphs
    that follow until the next marker or heading. `Why:` citation lines
    + `[ Click to enter... ]` placeholders are dropped (moved outside
    the edit zone so they don't count as evidence).
    """
    control_ref = _detect_arion_control_ref(md_text, file_name)
    if not control_ref:
        return md_text
    has_marker = bool(
        _ARION_MUST_LABEL_RE.search(md_text)
        or _ARION_SHOULD_LABEL_RE.search(md_text)
    )
    if not has_marker:
        return md_text

    # Strip Why lines + click placeholders from the whole doc (they carry
    # no tenant evidence and their content is fingerprint noise).
    md = _ARION_WHY_LABEL_RE.sub("", md_text)
    md = _ARION_CLICK_PLACEHOLDER_RE.sub("", md)

    # Split into lines and walk. For each ◆ label line: emit a
    # `<<MUST/SHOULD item:CTRL:slug>>` on its own line + open an edit
    # zone. Close the zone at the next ◆ label OR the next '## ' heading
    # OR the end of the section separator line (─── run).
    lines = md.split("\n")
    out: list[str] = []
    open_zone: str | None = None  # the item_id currently open
    for raw in lines:
        line = raw
        must_m = _ARION_MUST_LABEL_RE.match(line)
        should_m = _ARION_SHOULD_LABEL_RE.match(line)
        if must_m or should_m:
            # Close any prior zone
            if open_zone:
                out.append(f"<!-- EDIT-ZONE-END {open_zone} -->")
                open_zone = None
            slug = _slug_from_humanized((must_m or should_m).group(1))
            kind = "MUST" if must_m else "SHOULD"
            item_id = f"item:{control_ref}:{slug}"
            out.append(f"<<{kind} {item_id}>>")
            out.append(f"<!-- EDIT-ZONE-START {item_id} -->")
            open_zone = item_id
            continue
        # Section boundaries close the open zone
        stripped = line.strip()
        is_heading = stripped.startswith("#")
        is_hr      = bool(_re.fullmatch(r"[-─]{3,}", stripped))
        if open_zone and (is_heading or is_hr):
            out.append(f"<!-- EDIT-ZONE-END {open_zone} -->")
            open_zone = None
        out.append(line)
    if open_zone:
        out.append(f"<!-- EDIT-ZONE-END {open_zone} -->")

    return "\n".join(out)


_TABLE_SEPARATOR_RE = _re.compile(
    r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?\s*$"
)


def _synthesise_table_prose(md_text: str, max_chars: int = 12000) -> str:
    """Walk markdown table blocks in `md_text` and emit per-row prose
    the LLM can cite verbatim.

    Mammoth renders docx tables as pipe-delimited markdown rows. Each
    cell on its own is too short to clear the extractor's 40-char
    verbatim-quote bar — and the LLM that tries to synthesise across
    cells routinely hallucinates connections that grounding then
    rejects. Synthesising deterministically here gives the LLM real
    grounded prose to cite, captured into the doc text we trust.

    Output shape per table:
        Table N:
        Row 1: <Header1> = <cell1>; <Header2> = <cell2>; ...
        Row 2: ...

    Capped at `max_chars` total to bound the input growth on
    pathological documents. Bigger tables truncate gracefully.
    """
    if not md_text:
        return ""

    lines = md_text.split("\n")
    out_blocks: list[str] = []
    i = 0
    table_idx = 0
    total_chars = 0

    while i < len(lines) and total_chars < max_chars:
        line = lines[i]
        # A markdown table needs a separator on line i+1 (e.g. `| --- |`)
        if "|" in line and i + 1 < len(lines) and _TABLE_SEPARATOR_RE.match(lines[i + 1]):
            header_line = line
            data_lines: list[str] = []
            j = i + 2
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                data_lines.append(lines[j])
                j += 1

            headers = [c.strip() for c in header_line.strip().strip("|").split("|")]
            headers = [h for h in headers if h]

            row_proses: list[str] = []
            for row_num, dline in enumerate(data_lines, 1):
                cells = [c.strip() for c in dline.strip().strip("|").split("|")]
                if not any(c for c in cells):
                    continue
                pairs = []
                for h, c in zip(headers, cells):
                    if c:
                        pairs.append(f"{h} = {c}")
                if pairs:
                    row_proses.append(f"Row {row_num}: " + "; ".join(pairs) + ".")

            if row_proses:
                table_idx += 1
                block = f"Table {table_idx}:\n" + "\n".join(row_proses)
                if total_chars + len(block) > max_chars:
                    # Partial truncation — keep what fits, mark the rest
                    remaining = max_chars - total_chars
                    block = block[:remaining] + "\n[... table truncated ...]"
                    out_blocks.append(block)
                    total_chars = max_chars
                    break
                out_blocks.append(block)
                total_chars += len(block) + 2  # +2 for the joining "\n\n"

            i = j
        else:
            i += 1

    return "\n\n".join(out_blocks)


def _chunk_markdown_to_sections(
    md: str,
    target_chars: int = 20000,
) -> list[RawSection]:
    """Split markdown into synthetic sections targeting `target_chars`
    per chunk (~5K tokens). Used as a rescue when paragraph-walk misses
    table content.

    Mammoth often emits docx tables as a single long block with no
    `\\n\\n` separators, so we can't rely on paragraph boundaries
    alone — split on single newlines too, and if even a single 'line'
    is too long, hard-split it at byte boundaries. Sections carry no
    heading because mammoth's markdown doesn't reliably emit `#`
    syntax for docx-style headings.
    """
    if not md:
        return []

    # Prefer double-newline splits, then single-newline, then hard byte cuts
    units: list[str] = []
    for block in (md.split("\n\n") if "\n\n" in md else [md]):
        if len(block) <= target_chars:
            units.append(block)
            continue
        # Block too big — split on single newlines
        for line in block.split("\n"):
            if len(line) <= target_chars:
                units.append(line)
                continue
            # Single line too big — hard split
            for i in range(0, len(line), target_chars):
                units.append(line[i:i + target_chars])

    sections: list[RawSection] = []
    buf: list[str] = []
    buf_len = 0
    idx = 0

    def _flush():
        nonlocal buf, buf_len, idx
        if not buf:
            return
        sections.append(RawSection(
            section_id = f"md_chunk_{idx}",
            heading    = None,
            text       = "\n".join(buf),
            page_start = None,
            page_end   = None,
            level      = 0,
        ))
        idx += 1
        buf, buf_len = [], 0

    for u in units:
        if not u.strip():
            continue
        buf.append(u)
        buf_len += len(u) + 1
        if buf_len >= target_chars:
            _flush()
    _flush()
    return sections


# =============================================================================
# XLSX READER
# =============================================================================

# Fuzzy column name matching for compliance workbooks
_CONTROL_REF_ALIASES = [
    "control", "control_ref", "control ref", "iso ref", "clause",
    "control id", "controlid", "ref", "control number", "annex",
]
_FINDING_ALIASES = [
    "finding", "status", "compliance", "result", "assessment",
    "compliant", "gap status", "implementation",
]
_GAP_ALIASES = [
    "gap", "gap_description", "gap description", "comment", "notes",
    "evidence", "description", "observation", "detail",
]
_EVIDENCE_ALIASES = [
    "evidence", "evidence_text", "evidence text", "justification",
    "supporting evidence", "rationale",
]


def _fuzzy_col(headers: list[str], aliases: list[str]) -> Optional[int]:
    """Find column index by fuzzy name matching.
    Empty headers are skipped: `"" in alias` is always True and would
    otherwise match any alias against any blank column."""
    headers_lower = [h.lower().strip() for h in headers]
    for alias in aliases:
        for i, h in enumerate(headers_lower):
            if not h:
                continue
            if alias in h or h in alias:
                return i
    return None


def _is_compliance_workbook(headers: list[str]) -> bool:
    """Return True if the sheet looks like a compliance assessment workbook."""
    has_control = _fuzzy_col(headers, _CONTROL_REF_ALIASES) is not None
    has_finding = _fuzzy_col(headers, _FINDING_ALIASES) is not None
    return has_control and has_finding


# Standards used to validate the "control_ref" column actually contains
# parseable refs. Header detection only — the enricher still detects
# per-document standards from content keywords later.
_HEADER_DETECT_STANDARDS = ("ISO27001:2022", "GDPR:2016/679")


def _looks_like_ref(value) -> bool:
    """True if value normalizes to a control ref under any known standard."""
    from .ref_normalizer import normalize_ref  # local: avoid import cycle at module load
    if not value:
        return False
    s = str(value).strip()
    if not s:
        return False
    return any(normalize_ref(s, std) for std in _HEADER_DETECT_STANDARDS)


def _find_data_header(rows: list, max_scan: int = 30) -> Optional[tuple]:
    """Locate the real data-table header in a sheet, skipping document-metadata
    bands (Title/Owner/Revision History/etc.) that real-world ISMS workbooks
    place above the actual table.

    Returns (header_idx, ctrl_col, find_col, gap_col, evid_col) or None.

    Validation: the header row passes `_is_compliance_workbook`, has
    ctrl_col != find_col (rejects same-column degenerate fuzzy matches), and
    at least 2 of the next 8 non-empty values in ctrl_col normalize to a
    known control ref under any standard.
    """
    for idx, row in enumerate(rows[:max_scan]):
        if not any(c for c in row if c is not None):
            continue
        header = [str(c).strip() if c is not None else "" for c in row]
        ctrl_col = _fuzzy_col(header, _CONTROL_REF_ALIASES)
        find_col = _fuzzy_col(header, _FINDING_ALIASES)
        if ctrl_col is None or find_col is None or ctrl_col == find_col:
            continue
        examined = 0
        hits = 0
        for next_row in rows[idx + 1:idx + 1 + 30]:
            if examined >= 8:
                break
            if next_row is None or len(next_row) <= ctrl_col:
                continue
            val = next_row[ctrl_col]
            if val is None or not str(val).strip():
                continue
            examined += 1
            if _looks_like_ref(val):
                hits += 1
                if hits >= 2:
                    return (
                        idx,
                        ctrl_col,
                        find_col,
                        _fuzzy_col(header, _GAP_ALIASES),
                        _fuzzy_col(header, _EVIDENCE_ALIASES),
                    )
    return None


def _normalise_finding_value(val: str) -> str:
    """Map workbook finding values to canonical NC/OFI/Comply/N/A."""
    if not val:
        return "not_addressed"
    v = str(val).strip().lower()

    comply_terms  = ["comply", "compliant", "yes", "implemented", "done",
                     "complete", "full", "met", "pass", "✓", "green", "high"]
    ofi_terms     = ["ofi", "partial", "partly", "in progress", "improving",
                     "medium", "amber", "yellow", "opportunity"]
    nc_terms      = ["nc", "non-conform", "no", "not implemented", "fail",
                     "missing", "not met", "red", "critical", "not done"]
    na_terms      = ["n/a", "na", "not applicable", "out of scope", "excluded"]

    for term in comply_terms:
        if term in v:
            return "Comply"
    for term in ofi_terms:
        if term in v:
            return "OFI"
    for term in nc_terms:
        if term in v:
            return "NC"
    for term in na_terms:
        if term in v:
            return "N/A"
    return "not_addressed"


# Workbook meta sheets that aren't compliance data — skipped at read time so
# they don't show up as candidates for either path. Substring match on the
# lowercased sheet name. Surfaced 2026-06-23 on Arion's workbook: 5 of 38
# sheets are TOC / Documentation / Mapping / Instructions / Formulas, all
# spreadsheet metadata not posture evidence.
_META_SHEET_PATTERNS = (
    "table of contents", "toc",
    "documentation",
    "mapping",  # workbook's internal mapping reference, not compliance data
    "instructions", "definitions",
    "formulas", "formula",
    "key", "legend",
    "version history",  # different from compliance change-control logs
    "readme",
    "cover",  # cover sheets
    # Self-change-log for the workbook itself (Arion's "This Doc Chng Control"
    # — DOC005's change history with rows 0-12 self-metadata + 15+ change
    # entries on the workbook itself). The evidence shape doesn't match any
    # MUST leaf cleanly; better filtered than half-mapped. Added 2026-06-23.
    "doc chng control", "this doc chng",
    "doc change control", "doc change log",
    "chng control", "change control",
)


def _is_meta_sheet(sheet_name: str) -> bool:
    if not sheet_name:
        return False
    s = sheet_name.strip().lower()
    return any(p in s for p in _META_SHEET_PATTERNS)


def _filename_to_leaf_id(filename: str) -> Optional[str]:
    """Reverse-derive a leaf_id from a downloaded template's filename
    convention (see api_server._template_download_filename).

    The transform applied at download was:
        leaf_id  = "req:A.5.9:asset_inventory"
        filename = "A_5_9_asset_inventory.xlsx"  (req: stripped; : and . → _)

    Reversing without context is ambiguous (we can't tell where the
    control_ref ends and the slug begins), so we look up the templates
    table for any leaf_id whose computed download filename matches.

    Lazily imports the loader to avoid a hard dependency on the API
    helper from a reader module.
    """
    if not filename:
        return None
    base = filename.strip().lower()
    # Strip .xlsx and common tenant-added suffixes ("(1)", "_filled",
    # " - copy", trailing whitespace).
    base = base.removesuffix(".xlsx").removesuffix(".xlsm")
    base = _re.sub(r"\s*\(\d+\)\s*$", "", base)  # "name (1)"
    base = _re.sub(r"[_\s-]+(filled|copy|edited|v\d+)$", "", base)

    # Look up templates table to find a match. Lazy import + lazy
    # connection — most uploads have _arion_meta so this path is cold.
    try:
        from rag.posture_loader import build_pg_conn
    except Exception:
        return None
    try:
        conn = build_pg_conn()
        with conn.cursor() as cur:
            cur.execute("SELECT leaf_id FROM templates")
            for (leaf_id,) in cur.fetchall():
                if not leaf_id:
                    continue
                stripped = leaf_id[4:] if leaf_id.startswith("req:") else leaf_id
                expected = stripped.replace(":", "_").replace(".", "_").lower()
                if expected == base:
                    conn.close()
                    return leaf_id
        conn.close()
    except Exception as e:
        logger.warning(f"_filename_to_leaf_id: lookup failed: {e}")
    return None


def _derive_template_columns(leaf_id: str) -> tuple[list[str], list[str]]:
    """For filename-fallback identification: read the template body
    from the templates table and extract the canonical column /
    doc_field item_id lists.

    Returns (columns, doc_fields), both possibly empty.
    Used when `_arion_meta` is missing — the tenant deleted the
    hidden sheet — and we need an alternate source of truth.
    """
    try:
        from rag.posture_loader import build_pg_conn
        from rag.templates.xlsx_renderer import (
            _extract_table_columns, _extract_doc_level_must_ids,
        )
    except Exception:
        return ([], [])
    try:
        conn = build_pg_conn()
        with conn.cursor() as cur:
            cur.execute("SELECT body_md FROM templates WHERE leaf_id = %s",
                        (leaf_id,))
            row = cur.fetchone()
        conn.close()
        if not row:
            return ([], [])
        body = row[0] or ""
        columns    = _extract_table_columns(body, leaf_id) or []
        doc_fields = _extract_doc_level_must_ids(body, columns)
        return (columns, doc_fields)
    except Exception as e:
        logger.warning(f"_derive_template_columns: failed for {leaf_id}: {e}")
        return ([], [])


def _read_templated_xlsx_meta(wb, filename: str = "") -> Optional[dict]:
    """Detect whether this workbook was generated by our xlsx template
    renderer (rag/templates/xlsx_renderer.py).

    Two identification paths:

    1. **Hidden `_arion_meta` sheet** — primary, written at download
       time. Carries leaf_id + tenant_id + column_NN/doc_field_NN
       ordinals. Survives Excel save/load.

    2. **Filename-fallback** — when `_arion_meta` is missing (tenant
       deleted the hidden sheet), reverse-derive the leaf_id from
       the filename convention (`A_5_9_asset_inventory.xlsx`) by
       looking up the templates table for a matching expected
       filename. tenant_id is unknown via fallback; cross-tenant
       guard relies on upload-scope auth instead.

    In either path, also runs **column-count validation** — Register
    sheet width must match the canonical column_count from the meta
    (or, on fallback, from the template body). Mismatch → return None
    + log warning so the file falls through to the generic workbook
    lane (zero findings, safe default — avoids misbinding columns
    after tenant column-delete or reorder).

    Returns the parsed structure, or None when this isn't one of our
    templates (or it failed validation).
    """
    leaf_id: Optional[str]   = None
    tenant_id: str           = ""
    columns: list[str]       = []
    doc_fields: list[str]    = []
    source: str              = ""   # telemetry: "meta" | "filename_fallback"

    if "_arion_meta" in wb.sheetnames:
        source = "meta"
        try:
            m = wb["_arion_meta"]
            rows = list(m.iter_rows(values_only=True))
        except Exception:
            return None
        kv: dict[str, str] = {}
        for row in rows[1:]:  # skip header "key,value"
            if not row or len(row) < 2:
                continue
            k, v = row[0], row[1]
            if k is None:
                continue
            kv[str(k).strip()] = "" if v is None else str(v).strip()

        leaf_id   = kv.get("leaf_id")
        tenant_id = kv.get("tenant_id") or ""
        if not leaf_id:
            return None  # malformed meta — let generic path handle it

        # Reassemble ordered column + doc_field id lists from sparse keys
        ix = 0
        while True:
            v = kv.get(f"column_{ix:02d}")
            if not v:
                break
            columns.append(v)
            ix += 1
        ix = 0
        while True:
            v = kv.get(f"doc_field_{ix:02d}")
            if not v:
                break
            doc_fields.append(v)
            ix += 1
    else:
        # Filename-fallback: tenant deleted the hidden meta sheet (or
        # never had it). Try to recover leaf_id from the filename, then
        # re-derive column ordinals from the template body itself. The
        # uploading tenant's scope still applies via auth + RLS, so the
        # missing tenant_id field is acceptable here.
        recovered = _filename_to_leaf_id(filename)
        if not recovered:
            return None
        leaf_id = recovered
        columns, doc_fields = _derive_template_columns(leaf_id)
        if not columns and not doc_fields:
            return None  # could derive nothing — give up
        source = "filename_fallback"

    # Read Register sheet rows (data only — skip header + separator)
    # + column-count validation: if the Register header row width
    # doesn't match the canonical column_count, the tenant must have
    # added or removed columns. Bail out with a warning so the file
    # falls through to the generic workbook lane rather than misbinding
    # cells to wrong MUSTs.
    register_rows: list[dict] = []
    register_width: Optional[int] = None
    if "Register" in wb.sheetnames:
        ws = wb["Register"]
        rows_iter = list(ws.iter_rows(values_only=True))
        if rows_iter:
            # Header width — last non-None index + 1, since trailing
            # cells often come through as None even after edits.
            hdr = rows_iter[0]
            last_non_none = -1
            for i, cell in enumerate(hdr or ()):
                if cell is not None and str(cell).strip():
                    last_non_none = i
            register_width = last_non_none + 1 if last_non_none >= 0 else 0
        for row_ix, row in enumerate(rows_iter):
            if row_ix == 0:
                continue  # header
            payload: dict[int, str] = {}
            for col_ix, cell in enumerate(row):
                if cell is None:
                    continue
                text = str(cell).strip()
                if not text:
                    continue
                payload[col_ix] = text
            if payload:
                register_rows.append(payload)

    if columns and register_width is not None:
        if register_width != len(columns):
            logger.warning(
                f"templated_xlsx column-count mismatch on {filename!r}: "
                f"Register header has {register_width} non-empty cells but "
                f"expected {len(columns)} (leaf_id={leaf_id!r}, source={source!r}). "
                f"Possible tenant column-delete/add — falling through to "
                f"generic workbook lane to avoid misbinding."
            )
            return None

    # Read Document Fields sheet (hybrid templates only)
    doc_field_rows: list[dict] = []
    if "Document Fields" in wb.sheetnames:
        ws = wb["Document Fields"]
        for row_ix, row in enumerate(ws.iter_rows(values_only=True)):
            if row_ix == 0:
                continue  # header "Field | Standard text | Your content"
            label   = (str(row[0]).strip() if len(row) > 0 and row[0] is not None else "")
            stext   = (str(row[1]).strip() if len(row) > 1 and row[1] is not None else "")
            content = (str(row[2]).strip() if len(row) > 2 and row[2] is not None else "")
            doc_field_rows.append({
                "label":         label,
                "standard_text": stext,
                "your_content":  content,
            })

    return {
        "leaf_id":         leaf_id,
        "tenant_id":       tenant_id,
        "columns":         columns,
        "doc_fields":      doc_fields,
        "register_rows":   register_rows,
        "doc_field_rows":  doc_field_rows,
        "source":          source,
    }


def _read_xlsx(file_path: str, file_name: str) -> ParsedDocument:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl required: pip install openpyxl")

    wb       = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

    # ── Templated-upload detection ─────────────────────────────────────
    # If the workbook has our `_arion_meta` hidden sheet, this is a
    # round-trip upload of one of our generated templates. Stash the
    # parsed metadata + raw row data on the ParsedDocument so the
    # extractor can route through _extract_templated_xlsx without
    # re-parsing the file. See xlsx_renderer.py for the produced shape
    # and [[template-native-formats-xlsx-2026-06-26]] for the
    # foundation that made this possible.
    templated_meta = _read_templated_xlsx_meta(wb, filename=file_name)

    sections = []
    skipped_meta = []

    for sheet_name in wb.sheetnames:
        if _is_meta_sheet(sheet_name):
            skipped_meta.append(sheet_name)
            continue

        ws = wb[sheet_name]

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # Locate the real data-table header, skipping any document-metadata
        # band (Title/Owner/Revision History/etc.). Returns None if the sheet
        # is not a structured compliance table — those fall through to the
        # narrative branch below.
        hdr = _find_data_header(rows)

        if hdr is not None:
            header_idx, col_ctrl, col_find, col_gap, col_evid = hdr
            header_row = [
                str(c).strip() if c is not None else "" for c in rows[header_idx]
            ]

            structured_rows = []
            for row in rows[header_idx + 1:]:
                ctrl = row[col_ctrl] if col_ctrl is not None and len(row) > col_ctrl else None
                find = row[col_find] if col_find is not None and len(row) > col_find else None
                if not ctrl or not find:
                    continue
                structured_rows.append({
                    "control_ref":    str(ctrl).strip(),
                    "finding_raw":    str(find).strip(),
                    "finding":        _normalise_finding_value(str(find)),
                    "gap_description": str(row[col_gap]).strip() if col_gap is not None and len(row) > col_gap and row[col_gap] else "",
                    "evidence_text":   str(row[col_evid]).strip() if col_evid is not None and len(row) > col_evid and row[col_evid] else "",
                })

            # Represent as a single section with structured metadata
            text = f"COMPLIANCE WORKBOOK SHEET: {sheet_name}\n"
            text += "\n".join(
                f"{r['control_ref']} | {r['finding']} | {r['gap_description']}"
                for r in structured_rows
            )
            sections.append(RawSection(
                section_id = f"sheet_{sheet_name}",
                heading    = sheet_name,
                text       = text,
                page_start = None,
                page_end   = None,
                level      = 0,
                metadata   = {
                    "structured":    True,
                    "sheet_name":    sheet_name,
                    "column_map": {
                        "control_ref": header_row[col_ctrl] if col_ctrl is not None else None,
                        "finding":     header_row[col_find] if col_find is not None else None,
                        "gap":         header_row[col_gap]  if col_gap  is not None else None,
                    },
                    "rows": structured_rows,
                },
            ))
        else:
            # Narrative sheet — extract as text
            text_lines = []
            for row in rows:
                line = " | ".join(
                    str(c).strip() for c in row if c is not None and str(c).strip()
                )
                if line:
                    text_lines.append(line)
            if text_lines:
                sections.append(RawSection(
                    section_id = f"sheet_{sheet_name}",
                    heading    = sheet_name,
                    text       = "\n".join(text_lines),
                    page_start = None,
                    page_end   = None,
                    level      = 0,
                    metadata   = {"structured": False, "sheet_name": sheet_name},
                ))

    wb.close()

    doc = ParsedDocument(
        source_file   = file_path,
        file_type     = "xlsx",
        original_name = file_name,
        raw_sections  = sections,
        page_count    = 0,
    )
    if skipped_meta:
        doc.extraction_metrics["workbook_skipped_meta_sheets"] = ", ".join(skipped_meta)
        logger.info(
            f"_read_xlsx: skipped {len(skipped_meta)} meta sheet(s): "
            f"{', '.join(skipped_meta)}"
        )
    if templated_meta:
        doc.extraction_metrics["templated_xlsx_meta"] = templated_meta
        logger.info(
            f"_read_xlsx: templated upload detected — leaf_id="
            f"{templated_meta['leaf_id']}, "
            f"{len(templated_meta.get('register_rows', []))} register rows, "
            f"{len([r for r in templated_meta.get('doc_field_rows', []) if r.get('your_content')])} "
            f"doc_field rows with content"
        )
    return doc


# =============================================================================
# TXT / CSV READERS
# =============================================================================

def _read_txt(file_path: str, file_name: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    # Split at double newlines (paragraph boundaries)
    import re
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

    sections = []
    for i, para in enumerate(paragraphs):
        sections.append(RawSection(
            section_id = f"para_{i}",
            heading    = None,
            text       = para,
            page_start = None,
            page_end   = None,
            level      = 0,
        ))

    return ParsedDocument(
        source_file   = file_path,
        file_type     = "txt",
        original_name = file_name,
        raw_sections  = sections,
        page_count    = 0,
    )


def _read_csv(file_path: str, file_name: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8-sig", errors="replace") as f:
        reader = csv.DictReader(f)
        headers = reader.fieldnames or []
        rows    = list(reader)

    if _is_compliance_workbook(headers):
        # Treat like a structured XLSX sheet
        col_ctrl = _fuzzy_col(headers, _CONTROL_REF_ALIASES)
        col_find = _fuzzy_col(headers, _FINDING_ALIASES)
        col_gap  = _fuzzy_col(headers, _GAP_ALIASES)

        structured_rows = []
        for row in rows:
            ctrl = row.get(headers[col_ctrl], "") if col_ctrl is not None else ""
            find = row.get(headers[col_find], "") if col_find is not None else ""
            if not ctrl or not find:
                continue
            structured_rows.append({
                "control_ref":     ctrl.strip(),
                "finding":         _normalise_finding_value(find),
                "gap_description": row.get(headers[col_gap], "").strip() if col_gap is not None else "",
            })

        text = "COMPLIANCE CSV\n" + "\n".join(
            f"{r['control_ref']} | {r['finding']} | {r['gap_description']}"
            for r in structured_rows
        )
        section = RawSection(
            section_id = "csv_data",
            heading    = "CSV Data",
            text       = text,
            page_start = None,
            page_end   = None,
            level      = 0,
            metadata   = {"structured": True, "rows": structured_rows},
        )
    else:
        text = "\n".join(
            " | ".join(str(v) for v in row.values() if v)
            for row in rows
        )
        section = RawSection(
            section_id = "csv_data",
            heading    = "CSV Data",
            text       = text,
            page_start = None,
            page_end   = None,
            level      = 0,
            metadata   = {"structured": False},
        )

    return ParsedDocument(
        source_file   = file_path,
        file_type     = "csv",
        original_name = file_name,
        raw_sections  = [section],
        page_count    = 0,
    )


# =============================================================================
# HELPERS
# =============================================================================

def _detect_heading_level(heading: str) -> int:
    """Detect heading level from numbering pattern."""
    import re
    if not heading:
        return 0
    m = re.match(r'^(\d+)(\.\d+)*', heading)
    if m:
        dots = heading[:m.end()].count(".")
        return dots + 1
    if heading.isupper():
        return 1
    return 0
