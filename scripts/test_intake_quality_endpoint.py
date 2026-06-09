#!/usr/bin/env python3
"""
End-to-end smoke test for the intake quality-telemetry endpoint
(schema_v35, commit b74ece7).

Uploads a deliberately-bad docx — filename pattern that matches a
doc_mapping umbrella, body content that has no substantive
compliance evidence — and asserts that
GET /api/v1/admin/uploads/quality flags it RED with reason "0
findings from N scoped controls".

Useful for:
  - regression testing after schema_v35 changes (new drop bucket,
    new coverage signal)
  - validating the quality-flag thresholds after tuning
  - smoke-testing a fresh deploy

Run:
  PYTHONPATH=/data/arioncomply python3 scripts/test_intake_quality_endpoint.py

Exits 0 on success, 1 on any assertion / pipeline failure.
The DB rows and the /tmp stub file are cleaned up at the end.
"""
from __future__ import annotations

import argparse
import os
import sys
import time
from pathlib import Path

import psycopg2
import requests

# Default tenant: Arion. Override with --tenant-id for other tenants.
DEFAULT_TENANT_ID = "00000000-0000-0000-0000-000000000001"
DEFAULT_API_KEY   = "arion_dev_key_2026"
DEFAULT_BASE_URL  = "http://localhost:8080"
DEFAULT_DB_URL    = "postgresql://arioncomply:arioncomply2026@127.0.0.1/arioncomply_compliance"

STUB_FILENAME = "Risk_Management_Policy_smoketest.docx"
STUB_PATH     = Path("/tmp") / STUB_FILENAME


def _write_stub_docx(path: Path) -> None:
    """Create a docx that will trigger a doc_mappings filename match
    (Risk Management Policy umbrella → 5 candidate controls) but
    contains no substantive risk content the LLM can cite."""
    from docx import Document
    doc = Document()
    doc.add_heading("Risk Management Policy", level=1)
    doc.add_paragraph("This document is a placeholder. Content to be drafted in Q3.")
    doc.add_paragraph("Owner: TBD")
    doc.add_paragraph("Status: Draft")
    doc.add_paragraph("We will define our risk management approach soon.")
    doc.add_paragraph("See the SharePoint folder for related materials.")
    doc.save(str(path))


def _upload(base_url: str, api_key: str, path: Path) -> str:
    with path.open("rb") as f:
        r = requests.post(
            f"{base_url}/api/v1/documents/upload",
            headers={"X-API-Key": api_key},
            files={"file": (path.name, f)},
            timeout=30,
        )
    r.raise_for_status()
    return r.json()["upload_id"]


def _poll_until_done(base_url: str, api_key: str, upload_id: str, deadline_s: int = 120) -> dict:
    t_end = time.time() + deadline_s
    while time.time() < t_end:
        r = requests.get(
            f"{base_url}/api/v1/documents/{upload_id}/status",
            headers={"X-API-Key": api_key},
            timeout=10,
        )
        if r.status_code == 200:
            status = r.json()
            if status.get("status") in ("completed", "extracted", "failed"):
                return status
        time.sleep(2)
    raise TimeoutError(f"Upload {upload_id} did not complete within {deadline_s}s")


def _quality(base_url: str, api_key: str, flag: str | None = None) -> list[dict]:
    params = {"limit": 50}
    if flag:
        params["flag"] = flag
    r = requests.get(
        f"{base_url}/api/v1/admin/uploads/quality",
        headers={"X-API-Key": api_key},
        params=params,
        timeout=10,
    )
    r.raise_for_status()
    return r.json().get("uploads", [])


def _cleanup_db(db_url: str, tenant_id: str, filename: str) -> None:
    conn = psycopg2.connect(db_url)
    try:
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM intake_trace_log WHERE tenant_id = %s::uuid AND filename = %s",
                (tenant_id, filename),
            )
            cur.execute(
                "DELETE FROM document_uploads WHERE tenant_id = %s::uuid AND filename = %s",
                (tenant_id, filename),
            )
        conn.commit()
    finally:
        conn.close()


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--base-url",  default=DEFAULT_BASE_URL)
    ap.add_argument("--api-key",   default=DEFAULT_API_KEY)
    ap.add_argument("--tenant-id", default=DEFAULT_TENANT_ID)
    ap.add_argument("--db-url",    default=DEFAULT_DB_URL)
    ap.add_argument("--keep",      action="store_true",
                    help="Keep DB rows + tmp file after assertions (for inspection)")
    args = ap.parse_args()

    print(f"[1/5] Writing stub docx → {STUB_PATH}")
    _write_stub_docx(STUB_PATH)

    upload_id = None
    try:
        print(f"[2/5] Uploading via {args.base_url}/api/v1/documents/upload")
        upload_id = _upload(args.base_url, args.api_key, STUB_PATH)
        print(f"      upload_id={upload_id}")

        print(f"[3/5] Polling until pipeline completes (≤120s)")
        status = _poll_until_done(args.base_url, args.api_key, upload_id)
        print(f"      status={status.get('status')} findings={status.get('findings_count', 0)}")

        print(f"[4/5] Querying /api/v1/admin/uploads/quality?flag=red")
        red = _quality(args.base_url, args.api_key, flag="red")
        match = [u for u in red if u["upload_id"] == upload_id]
        if not match:
            all_recent = _quality(args.base_url, args.api_key)
            mine = [u for u in all_recent if u["upload_id"] == upload_id]
            actual = mine[0] if mine else None
            print(f"FAIL: stub not in red flag list. Actual: {actual}")
            return 1
        u = match[0]
        print(f"      [{u['quality_flag']}] {u['filename']} — {u['quality_reason']}")
        print(f"      cand={u['candidate_controls']} kept={u['findings_kept']} calls={u['llm_calls']}")

        print(f"[5/5] Assertions")
        assert u["quality_flag"] == "red", f"expected red, got {u['quality_flag']}"
        assert u["findings_kept"] == 0,    f"expected 0 findings, got {u['findings_kept']}"
        assert (u["candidate_controls"] or 0) >= 1, \
               f"expected ≥1 candidate (doc_mappings should match), got {u['candidate_controls']}"
        assert "0 findings" in u["quality_reason"], \
               f"unexpected reason: {u['quality_reason']}"
        print(f"      PASS")
        return 0

    except AssertionError as e:
        print(f"FAIL: {e}")
        return 1
    except Exception as e:
        print(f"FAIL: {type(e).__name__}: {e}")
        return 1
    finally:
        if not args.keep:
            print(f"[cleanup] Removing DB rows + {STUB_PATH}")
            try:
                _cleanup_db(args.db_url, args.tenant_id, STUB_FILENAME)
            except Exception as e:
                print(f"          db cleanup warning: {e}")
            try:
                STUB_PATH.unlink(missing_ok=True)
            except Exception as e:
                print(f"          file cleanup warning: {e}")


if __name__ == "__main__":
    sys.exit(main())
