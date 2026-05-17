"""
Integration test: confirm Stage 1 captures DOCX *table* content and Stage 4
persists the parsed markdown to document_text.

Pre-change behavior (the bug we are fixing in schema_v18):
  _read_docx walks only doc.paragraphs, dropping table cells entirely.
  A policy whose control evidence lives in a table → 0 findings.

Post-change behavior:
  _read_docx runs mammoth → markdown that includes tables. The extractor
  feeds markdown to the LLM, so table-only evidence yields findings.

This test builds a synthetic .docx in which the ONLY mention of A.5.30
(ICT readiness for business continuity) lives in a table cell — the
paragraphs around it are deliberately generic. It then runs the pipeline
end-to-end and asserts:

  1. findings_count >= 1
  2. at least one finding cites A.5.30 (table-only evidence reached the LLM)
  3. document_text row exists with markdown containing the table cell text

This case would have failed pre-change and passes post-change — matches the
feedback rule on adding regression coverage with each user-facing change.

Run:
  PYTHONPATH=/data/arioncomply python3 tests/test_intake_table_extraction.py
"""
from __future__ import annotations

import os
import sys
import uuid
from pathlib import Path

_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(_ROOT))

from dotenv import load_dotenv
load_dotenv(_ROOT / ".env")

import psycopg2
from docx import Document

from rag.intake.doc_pipeline import DocumentPipeline


TENANT_ID  = "00000000-0000-0000-0000-000000000001"   # Arion Networks
TABLE_ONLY_CONTROL = "A.5.30"
TABLE_ONLY_TEXT = (
    "ICT readiness for business continuity: RTO 4h, RPO 1h, validated "
    "via quarterly failover drills. Owner: ISMS Manager."
)


def _build_fixture(path: Path) -> None:
    """Synthesize a .docx where A.5.30 evidence lives ONLY in a table cell."""
    doc = Document()

    doc.add_heading("Business Continuity Policy (Test Fixture)", level=1)

    doc.add_heading("1. Purpose", level=2)
    doc.add_paragraph(
        "This policy describes the organization's general posture on "
        "operational resilience and the obligations of staff during "
        "service disruption events."
    )

    doc.add_heading("2. Scope", level=2)
    doc.add_paragraph(
        "This policy applies to all employees and contractors who operate "
        "production systems."
    )

    # The ONLY mention of A.5.30 lives in this table — pre-change the
    # paragraph walk drops it entirely, and the LLM never sees it.
    doc.add_heading("3. ICT Continuity Commitments", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Implementation"
    table.cell(1, 0).text = TABLE_ONLY_CONTROL
    table.cell(1, 1).text = TABLE_ONLY_TEXT

    doc.add_heading("4. Review", level=2)
    doc.add_paragraph("This policy is reviewed annually.")

    doc.save(path)


def main() -> int:
    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key:
        print("[FAIL] ANTHROPIC_API_KEY not set — cannot run extractor")
        return 1

    db_url = os.getenv("DATABASE_URL", "")
    if not db_url:
        print("[FAIL] DATABASE_URL not set")
        return 1

    upload_id = str(uuid.uuid4())
    fixture_dir = _ROOT / "tests" / "fixtures"
    fixture_dir.mkdir(parents=True, exist_ok=True)
    fixture_path = fixture_dir / f"_test_table_only_{upload_id[:8]}.docx"
    _build_fixture(fixture_path)

    # Pre-register the upload row so document_text's FK is satisfied.
    conn = psycopg2.connect(db_url)
    try:
        with conn.cursor() as cur:
            cur.execute("SET app.tenant_id = %s", (TENANT_ID,))
            cur.execute(
                """
                INSERT INTO document_uploads (
                    id, tenant_id, filename, storage_path,
                    extraction_status, sha256, byte_size
                ) VALUES (%s::uuid, %s::uuid, %s, %s, 'pending', %s, %s)
                """,
                (
                    upload_id,
                    TENANT_ID,
                    fixture_path.name,
                    str(fixture_path),
                    "test-fixture",
                    fixture_path.stat().st_size,
                ),
            )
        conn.commit()
    finally:
        conn.close()

    pipeline = DocumentPipeline(db_url=db_url, api_key=api_key, trace=True)
    result = pipeline.run(
        file_path         = str(fixture_path),
        tenant_id         = TENANT_ID,
        upload_id         = upload_id,
        original_filename = fixture_path.name,
    )

    print(f"[run] status={result.status} findings={result.findings_count}")
    print(f"[run] controls_assessed={result.controls_assessed}")

    # ── Assertion 1: at least one finding made it through ───────────────────
    if result.findings_count < 1:
        print(f"[FAIL] expected >=1 finding, got {result.findings_count}")
        _cleanup(upload_id, fixture_path)
        return 1

    # ── Assertion 2: the table-only control is among the findings ───────────
    matched = any(
        TABLE_ONLY_CONTROL in (ref or "")
        for ref in result.controls_assessed
    )
    if not matched:
        print(
            f"[FAIL] {TABLE_ONLY_CONTROL} not in controls_assessed — table "
            f"content did not reach the LLM. got={result.controls_assessed}"
        )
        _cleanup(upload_id, fixture_path)
        return 1

    # ── Assertion 3: document_text row exists and markdown captured table ──
    conn = psycopg2.connect(db_url)
    try:
        with conn.cursor() as cur:
            cur.execute("SET app.tenant_id = %s", (TENANT_ID,))
            cur.execute(
                """
                SELECT markdown, converter, byte_count, source_sha256
                FROM document_text
                WHERE upload_id = %s::uuid
                """,
                (upload_id,),
            )
            row = cur.fetchone()
    finally:
        conn.close()

    if row is None:
        print("[FAIL] no document_text row for this upload")
        _cleanup(upload_id, fixture_path)
        return 1

    md, converter, byte_count, src_sha = row
    # Mammoth escapes literal periods in markdown (A.5.30 → A\.5\.30), so
    # strip backslashes before substring checks. Also look for a unique
    # phrase from the table cell as a second signal that table content
    # really made it through.
    md_unescaped = md.replace("\\", "")
    if TABLE_ONLY_CONTROL not in md_unescaped:
        print(
            f"[FAIL] document_text.markdown missing {TABLE_ONLY_CONTROL} — "
            f"table content was dropped during conversion"
        )
        _cleanup(upload_id, fixture_path)
        return 1
    if "quarterly failover drills" not in md_unescaped:
        print(
            "[FAIL] document_text.markdown missing the table-cell evidence "
            "phrase — mammoth lost the table content"
        )
        _cleanup(upload_id, fixture_path)
        return 1

    if not converter.startswith("mammoth"):
        print(f"[FAIL] expected mammoth converter, got {converter!r}")
        _cleanup(upload_id, fixture_path)
        return 1

    print(
        f"[PASS] findings={result.findings_count}  "
        f"{TABLE_ONLY_CONTROL} matched  "
        f"md_bytes={byte_count}  converter={converter}  "
        f"src_sha={src_sha[:12]}..."
    )

    _cleanup(upload_id, fixture_path)
    return 0


def _cleanup(upload_id: str, fixture_path: Path) -> None:
    """Best-effort cleanup so repeated runs don't accumulate junk."""
    try:
        conn = psycopg2.connect(os.getenv("DATABASE_URL", ""))
        with conn.cursor() as cur:
            cur.execute("SET app.tenant_id = %s", (TENANT_ID,))
            # findings → client_documents.id; the pipeline links by storage_path
            cur.execute(
                """
                DELETE FROM document_findings
                WHERE document_id IN (
                    SELECT cd.id FROM client_documents cd
                    JOIN document_uploads du ON du.storage_path = cd.storage_path
                    WHERE du.id = %s::uuid
                )
                """,
                (upload_id,),
            )
            cur.execute(
                """
                DELETE FROM client_documents
                WHERE storage_path IN (
                    SELECT storage_path FROM document_uploads WHERE id = %s::uuid
                )
                """,
                (upload_id,),
            )
            # document_text + document_uploads cascade together via FK
            cur.execute("DELETE FROM document_uploads WHERE id = %s::uuid", (upload_id,))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"[cleanup] non-fatal: {e}")

    try:
        if fixture_path.exists():
            fixture_path.unlink()
    except Exception:
        pass


if __name__ == "__main__":
    sys.exit(main())
