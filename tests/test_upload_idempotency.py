"""
Integration test: confirm uploads are idempotent under schema_v19.

Two layers under test:

  Layer 1 — source-byte dedup
    Enforced by the partial unique index uniq_document_uploads_tenant_sha256.
    The upload endpoint pre-checks (tenant_id, sha256) before writing the row
    or the file. This test exercises the index directly by attempting two
    INSERTs with the same (tenant_id, sha256, status<>'duplicate') and
    expecting the second to fail.

  Layer 2 — normalized-markdown dedup
    Enforced inside DocumentPipeline.run after Stage 1. Two files with
    different bytes (e.g. DOCX vs PDF export of the same content) hash to
    the same markdown_sha256; the pipeline detects this BEFORE any LLM call
    and short-circuits with status='duplicate', dup_of_upload_id pointing
    at the canonical row, and the redundant file unlinked.

Pre-change (pre-v19) behaviour:
  Two uploads of the same file → two document_uploads rows, two extraction
  runs, two finding sets. No constraint, no pipeline dedup.

Post-change behaviour exercised here:
  Layer 1 INSERT raises IntegrityError on the second row.
  Layer 2 pipeline returns PipelineResult(status='duplicate') and marks the
  upload row accordingly without writing document_text / document_findings.

Run:
  PYTHONPATH=/data/arioncomply python3 tests/test_upload_idempotency.py
"""
from __future__ import annotations

import os
import sys
import hashlib
import uuid
from pathlib import Path
from typing import Optional

_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(_ROOT))

from dotenv import load_dotenv
load_dotenv(_ROOT / ".env")

import psycopg2
from psycopg2 import errors as pg_errors
from docx import Document

from rag.intake.doc_pipeline import DocumentPipeline


TENANT_ID = "00000000-0000-0000-0000-000000000001"   # Arion Networks


def _build_fixture(path: Path, run_marker: str) -> None:
    """Tiny DOCX that names exactly one ISO control. run_marker is included
    so each test run produces a unique sha256 / markdown_sha256 — this test
    builds the *same* content twice to exercise the dedup paths, but the
    marker keeps test runs isolated from each other."""
    doc = Document()
    doc.add_heading(f"Idempotency Fixture {run_marker}", level=1)
    doc.add_paragraph(
        "Access control reviews are performed quarterly by the ISMS owner. "
        "Findings feed back into the corrective action register. "
        f"Run: {run_marker}"
    )
    doc.add_heading("Control coverage", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "A.5.18"
    table.cell(1, 1).text = "Comply — reviews logged in the ISMS register."
    doc.save(path)


def _set_tenant(conn) -> None:
    with conn.cursor() as cur:
        cur.execute("SET app.tenant_id = %s", (TENANT_ID,))


def _delete_upload(conn, upload_id: str, filename: Optional[str] = None) -> None:
    """Best-effort cleanup of any rows a prior run may have left behind.
    Order matters: document_findings is keyed by client_documents.id, not
    upload_id, so we cascade through filename."""
    try:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM document_text WHERE upload_id=%s::uuid", (upload_id,))
            if filename:
                cur.execute(
                    "DELETE FROM document_findings "
                    " WHERE document_id IN (SELECT id FROM client_documents "
                    "                        WHERE tenant_id=%s::uuid AND filename=%s)",
                    (TENANT_ID, filename),
                )
                cur.execute(
                    "DELETE FROM client_documents "
                    " WHERE tenant_id=%s::uuid AND filename=%s",
                    (TENANT_ID, filename),
                )
            cur.execute("DELETE FROM document_uploads WHERE id=%s::uuid", (upload_id,))
        conn.commit()
    except Exception:
        conn.rollback()


def test_layer1_source_byte_dedup(db_url: str) -> tuple[bool, str]:
    """Two INSERTs with the same (tenant_id, sha256) and active status must
    violate the partial unique index."""
    upload_a = str(uuid.uuid4())
    upload_b = str(uuid.uuid4())
    # Random sha per run so a crashed prior run doesn't strand a row that
    # blocks the next run.
    shared_sha = hashlib.sha256(uuid.uuid4().bytes).hexdigest()

    conn = psycopg2.connect(db_url)
    try:
        _set_tenant(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO document_uploads (
                    id, tenant_id, filename, storage_path,
                    extraction_status, sha256, byte_size
                ) VALUES (%s::uuid, %s::uuid, %s, %s, 'pending', %s, %s)
                """,
                (upload_a, TENANT_ID, "_idem_layer1_a.docx",
                 "/tmp/_idem_layer1_a.docx", shared_sha, 100),
            )
        conn.commit()

        # Second INSERT must raise.
        raised = False
        try:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO document_uploads (
                        id, tenant_id, filename, storage_path,
                        extraction_status, sha256, byte_size
                    ) VALUES (%s::uuid, %s::uuid, %s, %s, 'pending', %s, %s)
                    """,
                    (upload_b, TENANT_ID, "_idem_layer1_b.docx",
                     "/tmp/_idem_layer1_b.docx", shared_sha, 100),
                )
            conn.commit()
        except pg_errors.UniqueViolation:
            conn.rollback()
            raised = True
        except Exception as e:
            conn.rollback()
            return False, f"unexpected exception type: {type(e).__name__}: {e}"

        if not raised:
            return False, "second INSERT with duplicate sha256 was not blocked"
        return True, "second INSERT with duplicate sha256 raised UniqueViolation"
    finally:
        _delete_upload(conn, upload_a, "_idem_layer1_a.docx")
        _delete_upload(conn, upload_b, "_idem_layer1_b.docx")
        conn.close()


def test_layer2_markdown_dedup(db_url: str, api_key: str) -> tuple[bool, str]:
    """Two pipeline runs over the same content with *different* sha256
    values (we override the byte hash to simulate two different exports
    of the same logical document) must result in:

      - run 1: status='completed', findings written, document_text row
      - run 2: status='duplicate', dup_of_upload_id = upload_id of run 1,
               no document_text row written, fixture file deleted
    """
    upload_a    = str(uuid.uuid4())
    upload_b    = str(uuid.uuid4())
    run_marker  = upload_a[:8]   # both fixtures use the same marker → same content → same markdown_sha256
    fixture_dir = _ROOT / "tests" / "fixtures"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    path_a = fixture_dir / f"_idem_layer2_a_{run_marker}.docx"
    path_b = fixture_dir / f"_idem_layer2_b_{run_marker}.docx"

    # Same content in both files (same run_marker → identical markdown).
    _build_fixture(path_a, run_marker)
    _build_fixture(path_b, run_marker)

    # Real sha256 of A; we synthesize a DIFFERENT sha256 for B so Layer 1
    # doesn't fire (Layer 1 is tested above).
    sha_a = hashlib.sha256(path_a.read_bytes()).hexdigest()
    sha_b = "f" * 64   # guaranteed different from sha_a

    conn = psycopg2.connect(db_url)
    try:
        _set_tenant(conn)
        # Pre-register both rows (the pipeline expects an existing row).
        with conn.cursor() as cur:
            for uid, fp, sha in [(upload_a, path_a, sha_a), (upload_b, path_b, sha_b)]:
                cur.execute(
                    """
                    INSERT INTO document_uploads (
                        id, tenant_id, filename, storage_path,
                        extraction_status, sha256, byte_size
                    ) VALUES (%s::uuid, %s::uuid, %s, %s, 'pending', %s, %s)
                    """,
                    (uid, TENANT_ID, fp.name, str(fp), sha, fp.stat().st_size),
                )
        conn.commit()

        pipeline = DocumentPipeline(db_url=db_url, api_key=api_key, trace=True)

        result_a = pipeline.run(
            file_path=str(path_a),
            tenant_id=TENANT_ID,
            upload_id=upload_a,
            original_filename=path_a.name,
        )
        if result_a.status not in ("extracted", "completed", "manual_review"):
            return False, f"run 1 unexpected status: {result_a.status} ({result_a.error})"

        result_b = pipeline.run(
            file_path=str(path_b),
            tenant_id=TENANT_ID,
            upload_id=upload_b,
            original_filename=path_b.name,
        )
        if result_b.status != "duplicate":
            return False, f"run 2 expected status='duplicate', got '{result_b.status}'"

        # Verify DB state.
        with conn.cursor() as cur:
            cur.execute(
                "SELECT extraction_status, dup_of_upload_id "
                "FROM document_uploads WHERE id = %s::uuid",
                (upload_b,),
            )
            status_b, dup_of_b = cur.fetchone()
            if status_b != "duplicate":
                return False, f"upload_b status in DB is '{status_b}', expected 'duplicate'"
            if str(dup_of_b) != upload_a:
                return False, f"dup_of_upload_id={dup_of_b}, expected {upload_a}"

            cur.execute(
                "SELECT COUNT(*) FROM document_text WHERE upload_id = %s::uuid",
                (upload_b,),
            )
            (text_rows_b,) = cur.fetchone()
            if text_rows_b != 0:
                return False, f"document_text rows for upload_b: {text_rows_b}, expected 0"

            # No client_documents row should have been created for upload_b
            # either — Stage 4 (where that write happens) never ran. The
            # registry links via filename, so check by that.
            cur.execute(
                "SELECT COUNT(*) FROM client_documents "
                " WHERE tenant_id = %s::uuid AND filename = %s",
                (TENANT_ID, path_b.name),
            )
            (cd_rows_b,) = cur.fetchone()
            if cd_rows_b != 0:
                return False, f"client_documents rows for upload_b: {cd_rows_b}, expected 0"

        if path_b.exists():
            return False, f"fixture file for upload_b still exists: {path_b}"

        return True, "Layer 2 pipeline dedup: status=duplicate, dup_of set, no text/findings, file unlinked"
    finally:
        _delete_upload(conn, upload_a, path_a.name)
        _delete_upload(conn, upload_b, path_b.name)
        conn.close()
        for p in (path_a, path_b):
            try:
                if p.exists():
                    p.unlink()
            except Exception:
                pass


def main() -> int:
    db_url = os.getenv("DATABASE_URL", "")
    if not db_url:
        print("[FAIL] DATABASE_URL not set")
        return 1
    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key:
        print("[FAIL] ANTHROPIC_API_KEY not set — Layer 2 needs the extractor")
        return 1

    print("─" * 70)
    print("  Upload idempotency — Stage 1 (schema_v19)")
    print("─" * 70)

    failures = 0

    ok, msg = test_layer1_source_byte_dedup(db_url)
    print(f"  [{'PASS' if ok else 'FAIL'}] Layer 1 source-byte dedup")
    print(f"         {msg}")
    if not ok:
        failures += 1

    ok, msg = test_layer2_markdown_dedup(db_url, api_key)
    print(f"  [{'PASS' if ok else 'FAIL'}] Layer 2 markdown-content dedup")
    print(f"         {msg}")
    if not ok:
        failures += 1

    print("─" * 70)
    print(f"  {2 - failures}/2 passed")
    return 0 if failures == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
